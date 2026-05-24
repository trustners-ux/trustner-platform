#!/usr/bin/env python3
"""
Generate MeraSIP Calculator Training Guide as a Word (.docx) document
with professional formatting, tables, and styling.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Heading styles
for i in range(1, 5):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.name = 'Calibri'
    h_style.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4A)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 1'].font.bold = True
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 2'].font.bold = True
doc.styles['Heading 3'].font.size = Pt(13)
doc.styles['Heading 3'].font.bold = True


# ── Helper Functions ──
def add_colored_heading(text, level=1, color=None):
    """Add a heading with optional custom color."""
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_body(text, bold=False, italic=False, color=None, size=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def add_selling_tip(text):
    """Add a selling tip in italic with special formatting."""
    p = doc.add_paragraph()
    run = p.add_run('Selling Tip: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x3E)
    run.font.size = Pt(10.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x55, 0x33)
    # Add light green background via shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0FFF0"/>')
    pPr.append(shading)
    return p

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0D2B4A"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = str(cell_text)
            for paragraph in row_cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        # Alternate row shading
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA"/>')
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacing after table
    return table

def add_bullet(text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_section_divider():
    """Add a thin line divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

def add_calculator_section(name, url, tag, what_it_does, when_to_use, inputs, what_client_sees, selling_tip):
    """Add a complete calculator section."""
    doc.add_heading(name, level=3)

    # URL and Tag line
    p = doc.add_paragraph()
    run = p.add_run('URL: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(url)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0xBF)
    p.add_run('    ')
    run = p.add_run(f'Tag: {tag}')
    run.font.size = Pt(10)
    run.italic = True

    # What it does
    p = doc.add_paragraph()
    run = p.add_run('What it does: ')
    run.bold = True
    p.add_run(what_it_does)

    # When to use
    p = doc.add_paragraph()
    run = p.add_run('When to use it:')
    run.bold = True
    for item in when_to_use:
        add_bullet(item)

    # Inputs table
    if inputs:
        p = doc.add_paragraph()
        run = p.add_run('What the client enters:')
        run.bold = True
        add_table(['Input', 'Range / Options'], inputs)

    # What client sees
    p = doc.add_paragraph()
    run = p.add_run('What the client sees:')
    run.bold = True
    for item in what_client_sees:
        add_bullet(item)

    # Selling tip
    add_selling_tip(selling_tip)
    doc.add_paragraph()  # spacing


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

# Add some spacing at top
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MeraSIP')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x0B, 0x6E, 0xBF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Calculator Training Guide')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4A)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('For Trustner Team, POSPs & Sub-Brokers')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('30 Interactive Financial Calculators')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x3E)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Wealth & SIP  |  Loan  |  Tax  |  Insurance  |  Life Decisions')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Platform: www.merasip.com/calculators')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x0B, 0x6E, 0xBF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Trustner Asset Services Pvt. Ltd.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0  |  March 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1. Introduction',
    '2. How to Use the Calculators',
    '3. Category A: Wealth & SIP Calculators (15)',
    '4. Category B: Loan Calculators (4)',
    '5. Category C: Tax Calculators (3)',
    '6. Category D: Insurance Calculators (3)',
    '7. Category E: Life Decision Calculators (5)',
    '8. Calculator-to-Product Mapping',
    '9. Client Conversation Scripts',
    '10. Common Client Scenarios & Objection Handling',
    '11. Tips for POSPs & Sub-Brokers',
    '12. Quick Reference: All 30 Calculator URLs',
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4A)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)

doc.add_heading('What is MeraSIP?', level=2)
add_body('MeraSIP (www.merasip.com) is Trustner\'s digital investment education and planning platform. It provides 30 free interactive financial calculators that help investors make informed decisions about SIPs, loans, taxes, insurance, and major life financial choices.')

doc.add_heading('Why Should You Use These Calculators?', level=2)
benefits = [
    ('Build Trust: ', 'Show clients real numbers, not just promises'),
    ('Visual Impact: ', 'Every calculator has charts and graphs that make concepts easy to understand'),
    ('PDF Reports: ', 'Every calculator lets you download a PDF report that you can share with the client'),
    ('Free to Use: ', 'No login required, completely free for clients'),
    ('Mobile-Friendly: ', 'Works on phones, tablets, and computers'),
    ('Real-Time: ', 'All results update instantly as inputs change'),
]
for prefix, text in benefits:
    add_bullet(text, bold_prefix=prefix)

doc.add_heading('How This Guide is Organized', level=2)
add_body('Each calculator is explained with:')
items = [
    'What it does (one-line summary)',
    'When to use it (client situations)',
    'What the client enters (input fields)',
    'What the client sees (results and charts)',
    'Selling conversation tip (how to convert interest into action)',
]
for item in items:
    add_bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 2: HOW TO USE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('2. How to Use the Calculators', level=1)

doc.add_heading('Accessing the Calculators', level=2)
steps = [
    'Go to www.merasip.com/calculators',
    'You will see all 30 calculators organized into 5 categories with quick-jump navigation pills at the top',
    'Click any calculator card to open it',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'{step}')

doc.add_heading('Common Features (All 30 Calculators)', level=2)
features = [
    ('Left Panel: ', 'Input fields where you enter numbers (sliders, buttons, or type directly)'),
    ('Right Panel: ', 'Results, charts, and tables update in real-time'),
    ('Download PDF: ', 'Every calculator has a "Download PDF" button to save a report'),
    ('Back Button: ', 'Click the arrow at the top-left to return to the main page'),
    ('Disclaimer: ', 'All calculators show SEBI/AMFI/IRDAI disclaimers at the bottom'),
]
for prefix, text in features:
    add_bullet(text, bold_prefix=prefix)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Pro Tip for POSPs: ')
run.bold = True
run.font.color.rgb = RGBColor(0x0B, 0x6E, 0xBF)
p.add_run('When sitting with a client, enter THEIR actual numbers in the calculator. The personal touch of seeing their own income, expenses, and goals reflected in the charts creates much stronger engagement than showing generic examples.')
pPr = p._p.get_or_add_pPr()
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FF"/>')
pPr.append(shading)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 3: WEALTH & SIP CALCULATORS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. Category A: Wealth & SIP Calculators', level=1)
add_body('Navigate to: www.merasip.com/calculators#wealth', color=RGBColor(0x0B, 0x6E, 0xBF), size=11)
add_body('15 Calculators | These are your core investment planning tools. Use these daily with prospective and existing clients.', bold=True)

# 3.1 SIP
add_calculator_section(
    '3.1 SIP Future Value Calculator',
    'merasip.com/calculators/sip', 'Most Popular',
    'Shows how much a fixed monthly SIP will grow over time with the power of compounding.',
    ['First meeting with a new investor who asks "What will my SIP become?"',
     'When a client wants to understand compounding',
     'Quick projection for any SIP amount'],
    [['Monthly SIP Amount', 'Rs 500 - 5,00,000'],
     ['Expected Annual Return', '1% - 30%'],
     ['Investment Duration', '1 - 40 years']],
    ['Total Invested vs Total Value (the "wealth created" gap)',
     'Wealth gain amount and percentage',
     'Area chart showing growth curve over time',
     'Year-by-year breakdown table'],
    '"See how Rs 10,000/month becomes Rs 23 lakhs in 10 years? And you only invested Rs 12 lakhs. The remaining Rs 11 lakhs is the magic of compounding. Imagine what happens if we continue for 20 years..." Then increase the years to 20 and watch the client\'s expression change.'
)

# 3.2 Step-Up SIP
add_calculator_section(
    '3.2 Step-Up SIP Calculator',
    'merasip.com/calculators/step-up-sip', 'Recommended',
    'Shows the massive difference when you increase your SIP by a fixed percentage every year (e.g., 10% annual step-up).',
    ['When a client says "I\'ll increase my SIP when my salary increases"',
     'Salary increment discussions',
     'Showing the power of incremental discipline'],
    [['Starting Monthly SIP', 'Rs 500 - 5,00,000'],
     ['Annual Step-Up Rate', '1% - 50%'],
     ['Expected Return', '1% - 30%'],
     ['Duration', '1 - 40 years']],
    ['Side-by-side comparison: Regular SIP vs Step-Up SIP',
     'The "extra wealth" created by stepping up',
     'Chart showing both growth curves diverging over time'],
    '"If your salary goes up 10% every year, just increase your SIP by 10% too. Look at the difference -- Rs 45 lakhs with step-up vs Rs 23 lakhs without. That\'s almost double, just by matching your SIP to your salary growth."'
)

# 3.3 Goal-Based
add_calculator_section(
    '3.3 Goal-Based SIP Calculator',
    'merasip.com/calculators/goal-based', 'Goal Planning',
    'Reverse-calculates -- tells you exactly how much monthly SIP you need to reach a specific financial goal.',
    ['Client has a specific goal (child\'s education, marriage, house, retirement)',
     '"How much do I need to invest for [goal]?"',
     'Financial planning conversations'],
    [['Target Amount', 'Goal amount in rupees'],
     ['Time Horizon', 'Years to achieve goal'],
     ['Expected Return', 'Annual return %']],
    ['Required Monthly SIP amount',
     '3-scenario analysis (Conservative, Moderate, Aggressive returns)',
     'Growth projection chart'],
    '"Your daughter\'s engineering will cost Rs 25 lakhs in 15 years. You only need Rs 4,200 per month to get there. That\'s less than what many families spend on eating out. Shall we start this SIP today?"'
)

# 3.4 Inflation-Adjusted
add_calculator_section(
    '3.4 Inflation-Adjusted SIP Calculator',
    'merasip.com/calculators/inflation-adjusted', 'Essential',
    'Shows the REAL value of your SIP after accounting for inflation -- because Rs 1 crore in 20 years won\'t buy what Rs 1 crore buys today.',
    ['When a client thinks "Rs 1 crore is enough for retirement"',
     'Reality check conversations',
     'Explaining why higher SIP amounts are needed'],
    [['Monthly SIP', 'Amount'],
     ['Expected Return', '% p.a.'],
     ['Duration', 'Years'],
     ['Inflation Rate', '% (typically 6-7%)']],
    ['Nominal vs Real (inflation-adjusted) corpus',
     'How much purchasing power is lost to inflation',
     'Two-line chart showing the gap between nominal and real value'],
    '"Your SIP will grow to Rs 1 crore in 20 years. But with 6% inflation, that Rs 1 crore will only buy what Rs 31 lakhs buys today. That\'s why we need to plan for a higher target."'
)

# 3.5 Retirement
add_calculator_section(
    '3.5 Retirement SIP Planner',
    'merasip.com/calculators/retirement', 'Life Planning',
    'Calculates the monthly SIP needed to build your retirement corpus, considering current age, retirement age, monthly expenses, and inflation.',
    ['Retirement planning discussions',
     'Clients in their 30s-40s who haven\'t started planning',
     '"How much do I need for retirement?"'],
    [['Current Age', '18-65'],
     ['Retirement Age', '45-70'],
     ['Monthly Expenses', 'Current monthly spending'],
     ['Inflation Rate', '5-10%'],
     ['Expected Return', 'Pre and post-retirement'],
     ['Existing Savings', 'Current investments']],
    ['Required Retirement Corpus',
     'Monthly SIP needed',
     'Gap between current savings and target',
     'Year-by-year projection'],
    '"You\'re 35 and want to retire at 60. Your monthly expenses of Rs 50,000 will become Rs 2.15 lakhs by then due to inflation. You need Rs 4.5 crore corpus. Starting a SIP of Rs 28,000/month today will get you there."'
)

# 3.6 SWP
add_calculator_section(
    '3.6 SWP Calculator (Systematic Withdrawal Plan)',
    'merasip.com/calculators/swp', 'Post-Retirement',
    'Plans systematic withdrawals from a corpus -- how much you can withdraw monthly and how long the money will last.',
    ['Post-retirement planning',
     'Clients who have a lump sum and want regular income',
     '"How long will my Rs 50 lakhs last?"'],
    [['Total Corpus', 'Invested amount'],
     ['Monthly Withdrawal', 'Desired income'],
     ['Expected Return', 'Post-retirement return'],
     ['Duration', 'Withdrawal period']],
    ['Total withdrawn over the period',
     'Remaining corpus at the end',
     'Whether the money runs out (and when)',
     'Month-by-month depletion chart'],
    '"With Rs 50 lakhs invested at 8% return, you can withdraw Rs 40,000/month for 20 years. And you\'ll still have Rs 12 lakhs left at the end. SWP gives you pension-like income from mutual funds."'
)

# 3.7 Lumpsum
add_calculator_section(
    '3.7 Lumpsum Investment Calculator',
    'merasip.com/calculators/lumpsum', 'Investment',
    'Projects the growth of a one-time lump sum investment over time.',
    ['Client receives a bonus, inheritance, or property sale proceeds',
     'Comparing lumpsum vs SIP',
     '"I have Rs 10 lakhs to invest, what will it become?"'],
    [['Lumpsum Amount', 'Investment amount'],
     ['Expected Return', 'Annual return %'],
     ['Duration', 'Investment period']],
    ['Future value of the investment',
     'Total returns earned',
     'Growth curve chart'],
    '"Your Rs 10 lakh bonus invested today at 12% will become Rs 31 lakhs in 10 years. If you keep it in a savings account at 3.5%, it\'ll only be Rs 14 lakhs. That\'s Rs 17 lakhs of opportunity lost."'
)

# 3.8 SIP vs Lumpsum
add_calculator_section(
    '3.8 SIP vs Lumpsum Comparison',
    'merasip.com/calculators/sip-vs-lumpsum', 'Strategy',
    'Side-by-side comparison of investing the same total amount as a monthly SIP vs one-time lumpsum.',
    ['Client asks "Should I invest all at once or spread it out?"',
     'Market timing discussions',
     'Risk comparison conversations'],
    [['Total Investment Amount', 'Total to invest'],
     ['Expected Return', 'Annual return %'],
     ['Duration', 'Investment period']],
    ['SIP final value vs Lumpsum final value',
     'Which approach gives higher returns',
     'Detailed comparison chart'],
    '"Mathematically, lumpsum gives higher returns because your entire money works from day one. But SIP gives you rupee cost averaging -- you buy more units when markets are low. For most salaried investors, SIP is the practical and less stressful choice."'
)

# 3.9 Duration Optimizer
add_calculator_section(
    '3.9 SIP Duration Optimizer',
    'merasip.com/calculators/duration-optimizer', 'Optimizer',
    'Finds how many years it takes to reach a target corpus with a given monthly SIP and return rate.',
    ['"How long will it take to reach Rs 1 crore?"',
     'Setting realistic timelines for goals',
     'Adjusting SIP amount vs duration trade-offs'],
    [['Target Amount', 'Goal corpus'],
     ['Monthly SIP', 'Current SIP amount'],
     ['Expected Return', 'Annual return %']],
    ['Exact years and months to reach target',
     'Progress milestones along the way',
     'Growth trajectory chart'],
    '"At Rs 15,000/month with 12% returns, you\'ll reach Rs 1 crore in about 16 years. Want to reach it in 12 years? You\'d need Rs 25,000/month. Let\'s find the right balance for you."'
)

# 3.10 Market Correction
add_calculator_section(
    '3.10 Market Correction Impact Calculator',
    'merasip.com/calculators/correction-impact', 'Simulation',
    'Simulates a market crash during your SIP tenure and shows how corrections affect your portfolio long-term.',
    ['Client is worried about market crashes',
     'During actual market corrections (great reassurance tool)',
     'Explaining why SIPs work well in volatile markets'],
    [['Monthly SIP', 'SIP amount'],
     ['Expected Return', 'Annual return %'],
     ['Total Duration', 'Investment years'],
     ['Correction Year', 'When crash happens'],
     ['Correction Magnitude', 'How much market falls (%)'],
     ['Recovery Period', 'Years to recover']],
    ['Normal portfolio vs corrected portfolio comparison',
     'Impact amount and percentage',
     'Severity analysis table (10%, 20%, 30%, 40%, 50% corrections)',
     'Recovery trajectory chart'],
    '"Even if there\'s a 30% crash in year 5, your 20-year SIP portfolio only reduces by 8%. And because you bought more units during the crash, the recovery is faster. Market corrections are actually SIP investors\' best friends."'
)

# 3.11 Life-Stage
add_calculator_section(
    '3.11 Life-Stage Planner',
    'merasip.com/calculators/life-stage', 'Life Planning',
    'Plans your complete investment lifecycle through 3 phases: Invest (accumulate), Grow (let it compound), Withdraw (enjoy).',
    ['Comprehensive financial planning',
     'Showing the full journey from investing to retirement income',
     'Clients who want a "complete picture"'],
    [['Phase 1 - Monthly SIP', 'Investment amount'],
     ['Phase 1 - Duration', 'Accumulation years'],
     ['Phase 2 - Duration', 'Compounding years (no SIP, no withdrawal)'],
     ['Phase 3 - Monthly Withdrawal', 'Retirement income needed'],
     ['Phase 3 - Duration', 'Withdrawal years'],
     ['Return Rates', 'For each phase']],
    ['3-phase color-coded chart (green invest, teal grow, amber withdraw)',
     'Peak corpus value',
     'Total invested vs total withdrawn',
     'Wealth multiplier (how many times your investment grew)'],
    '"Invest Rs 15,000/month for 15 years, let it grow untouched for 5 years, then withdraw Rs 50,000/month for 20 years. Your total investment: Rs 27 lakhs. Total withdrawn: Rs 1.2 crores. That\'s the power of the 3-phase approach."'
)

# 3.12 Lifeline
add_calculator_section(
    '3.12 Lifeline Financial Planner',
    'merasip.com/calculators/lifeline', 'CFP-Style',
    'The most advanced calculator -- lets you add multiple financial events (SIPs, lump sums, withdrawals) at any point in your financial journey. Like what a Certified Financial Planner does.',
    ['Complex financial planning with multiple goals',
     'HNI clients with varied cash flows',
     'When simple calculators aren\'t enough'],
    [['Starting Corpus', 'Existing investments'],
     ['Return Rate', 'Expected annual return'],
     ['Planning Horizon', 'Total years to plan'],
     ['Life Events', 'Add up to 15 events: SIP, Lump Sum, SWP, Withdrawal']],
    ['Complete financial lifeline on one chart',
     'All events marked with reference lines',
     'Year-by-year corpus progression',
     'Net invested vs corpus value'],
    '"Let me create your complete financial plan right here. We\'ll add your current SIP, the bonus you\'ll invest in year 3, your child\'s education withdrawal in year 12, and your retirement SWP from year 20. One chart, your entire financial life."'
)

# 3.13 Daily SIP
add_calculator_section(
    '3.13 Daily SIP Calculator',
    'merasip.com/calculators/daily-sip', 'Micro-Invest',
    'Calculates returns on daily investments -- for apps that allow daily SIPs.',
    ['Clients interested in micro-investing',
     '"Even Rs 100/day can build wealth"',
     'Comparing daily vs monthly SIP'],
    [['Daily SIP Amount', 'Amount per day'],
     ['Expected Return', 'Annual return %'],
     ['Duration', 'Investment years'],
     ['Day Type', 'Calendar days or Working days']],
    ['Total value of daily investments',
     'Comparison with equivalent monthly SIP',
     'Growth chart'],
    '"Rs 300 per day -- that\'s the cost of a coffee and snack. In 15 years at 12% returns, it becomes Rs 45 lakhs. Small daily habits create massive wealth."'
)

# 3.14 FIRE
add_calculator_section(
    '3.14 FIRE Calculator (Financial Independence, Retire Early)',
    'merasip.com/calculators/fire', 'FIRE',
    'Calculates your Financial Independence number -- the corpus at which your investments generate enough income to cover all expenses forever, so you can retire early.',
    ['Young professionals (25-40) interested in early retirement',
     '"How much do I need to never work again?"',
     'High-income clients wanting to accelerate wealth'],
    [['FIRE Type', 'Regular (1x), Lean (0.7x), Fat (1.5x)'],
     ['Current Age', '18-60'],
     ['Target Retirement Age', '25-70'],
     ['Monthly Income & Expenses', 'Current amounts'],
     ['Current Savings', 'Existing investments'],
     ['Monthly Investment', 'SIP amount'],
     ['Return Rate', 'Pre and post-retirement'],
     ['Inflation Rate', '3-12%'],
     ['Safe Withdrawal Rate', '2-6% (the 4% rule)']],
    ['FIRE Number (the magic corpus)',
     'Years to FIRE with progress bar',
     'Coast FIRE Number (amount needed now if you stop investing)',
     'Two-phase chart: Accumulation (green) + Withdrawal (red)',
     'How long the money lasts'],
    '"Your FIRE number is Rs 3.2 crores. At your current savings rate, you\'ll reach it by age 48 -- that\'s retiring 12 years early. Every Rs 5,000 increase in monthly SIP brings your FIRE date 2 years closer."'
)

# 3.15 Cost of Delay
add_calculator_section(
    '3.15 Cost of Delay Calculator',
    'merasip.com/calculators/delay-cost', 'Eye-Opener',
    'The most powerful motivational calculator -- shows exactly how much wealth is destroyed by procrastinating. Compares scenarios of delaying SIP by 0, 1, 3, 5, and 10 years.',
    ['Client who says "I\'ll start next year" or "Let me think about it"',
     'Urgency creation in sales conversations',
     'Any first meeting to drive action'],
    [['Monthly SIP', 'SIP amount'],
     ['Expected Return', 'Annual return %'],
     ['Investment Horizon', 'Total years'],
     ['Delay Scenarios', '0, 1, 3, 5, 10 years (pre-built buttons)']],
    ['5 scenarios computed simultaneously',
     'Multi-line chart showing how each delay reduces final corpus',
     'Shock Value Card: Cost Per Day of Delay',
     'Required Catch-Up SIP (the higher SIP needed to match)'],
    '"Delaying your Rs 10,000 SIP by just 3 years costs you Rs 12 lakhs. That\'s Rs 1,100 per day of delay. And to catch up, you\'d need Rs 15,500/month instead of Rs 10,000. The best time to start was yesterday. The second best time is today."'
)

p = doc.add_paragraph()
run = p.add_run('This is your #1 closing tool. Use it in every meeting.')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.font.size = Pt(12)
pPr = p._p.get_or_add_pPr()
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF0F0"/>')
pPr.append(shading)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 4: LOAN CALCULATORS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('4. Category B: Loan Calculators', level=1)
add_body('Navigate to: www.merasip.com/calculators#loan', color=RGBColor(0x0B, 0x6E, 0xBF), size=11)
add_body('4 Calculators | Help clients make smarter borrowing decisions and save on interest costs.', bold=True)

# 4.1 EMI
add_calculator_section(
    '4.1 EMI Calculator',
    'merasip.com/calculators/emi', 'Popular',
    'Calculates monthly EMI for any loan type with full year-by-year amortization schedule.',
    ['Client taking a home/car/personal/education loan',
     'Comparing loan offers from different banks',
     'Understanding interest vs principal split'],
    [['Loan Amount', 'Rs 1L - 10 Cr'],
     ['Interest Rate', '1% - 25%'],
     ['Loan Tenure', '1 - 30 years'],
     ['Loan Type', 'Home / Car / Personal / Education (preset buttons)']],
    ['Monthly EMI amount',
     'Total Interest Paid and Total Amount Paid',
     'Interest-to-Principal Ratio',
     'Area chart: Principal vs Interest per year',
     'Pie chart: Principal vs Interest split',
     'Full year-by-year amortization table'],
    '"For a Rs 50 lakh home loan at 8.5% for 20 years, your EMI is Rs 43,391. But you\'ll pay Rs 54 lakhs in interest -- more than the loan itself! Now let me show you how prepayment can save lakhs..." (then switch to Prepayment Calculator)'
)

# 4.2 Loan Prepayment
add_calculator_section(
    '4.2 Loan Prepayment Calculator',
    'merasip.com/calculators/loan-prepayment', 'Save Interest',
    'Shows how much interest and tenure you save by making extra payments towards your loan.',
    ['Client with existing loan who gets a bonus/increment',
     '"Should I prepay my loan or invest?"',
     'Showing the power of small prepayments'],
    [['Loan Amount', 'Rs 1L - 10 Cr'],
     ['Interest Rate', '1% - 25%'],
     ['Loan Tenure', '1 - 30 years'],
     ['Prepayment Amount', 'Rs 10K - 10 Cr'],
     ['Prepayment Type', 'One-Time or Every Year'],
     ['Start Year', '1 to tenure']],
    ['Interest Saved (hero metric with gradient card)',
     'Tenure Reduction (years and months saved)',
     'Return on Prepayment (%)',
     'Balance comparison chart: With vs Without Prepayment',
     'Amortization table with "Prepaid" badges'],
    '"Prepaying Rs 2 lakhs every year on your Rs 50 lakh home loan saves you Rs 18 lakhs in interest and reduces tenure by 6 years. But if you invest that Rs 2 lakhs in SIP instead at 12% return, you\'ll earn Rs 28 lakhs more. Let me show you..."'
)

# 4.3 Car Loan vs Cash
add_calculator_section(
    '4.3 Car Loan vs Cash Calculator',
    'merasip.com/calculators/car-loan-vs-cash', 'Decision',
    'Compares whether it is better to take a car loan (and invest the cash) or pay cash outright. Includes opportunity cost analysis.',
    ['Client buying a car and has cash available',
     '"Should I take a car loan?"',
     'Teaching opportunity cost of cash payments'],
    [['Car Price', 'Rs 3L - 1 Cr'],
     ['Down Payment', '0 - 50%'],
     ['Loan Interest Rate', '5% - 18%'],
     ['Loan Tenure', '1 - 7 years'],
     ['Investment Return Rate', '6% - 20%']],
    ['Verdict: "Take the Loan" or "Pay Cash"',
     'Advantage amount and Break-even return rate',
     'Wealth trajectory comparison chart',
     'Detailed comparison table'],
    '"If the car costs Rs 8 lakhs and you can earn 12% on investments while the car loan is at 9%, it\'s actually smarter to take the loan and invest the Rs 8 lakhs. Over 5 years, you come out Rs 1.2 lakhs ahead."'
)

# 4.4 Home Affordability
add_calculator_section(
    '4.4 Home Affordability Calculator',
    'merasip.com/calculators/home-affordability', 'Home Buying',
    'Tells you exactly how much house you can afford based on your income, using 3-tier FOIR (Fixed Obligation to Income Ratio) analysis.',
    ['Client planning to buy a home',
     '"What budget should I look at?"',
     'Before home loan application'],
    [['Monthly Gross Income', 'Rs 20K - 50L'],
     ['Existing EMIs', 'Rs 0 - 5L'],
     ['Loan Interest Rate', '5% - 15%'],
     ['Loan Tenure', '5 - 30 years'],
     ['Down Payment', '5% - 50%']],
    ['3 Affordability Tiers: Comfortable (35% FOIR, Green), Manageable (45%, Amber), Stretched (55%, Red)',
     'Each tier shows: Max Property Price, Max EMI, Loan Amount, Down Payment',
     'Income Allocation pie chart',
     'Interest Rate Sensitivity chart',
     'Full affordability matrix (3 FOIR x 5 interest rates)'],
    '"With your Rs 1.5 lakh monthly income and Rs 15K existing EMI, you can comfortably afford a Rs 65 lakh home. You COULD stretch to Rs 90 lakhs, but that means 55% of income goes to EMIs. I\'d recommend staying in the green zone."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 5: TAX CALCULATORS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('5. Category C: Tax Calculators', level=1)
add_body('Navigate to: www.merasip.com/calculators#tax', color=RGBColor(0x0B, 0x6E, 0xBF), size=11)
add_body('3 Calculators | Help clients understand their tax liability and save more through proper planning.', bold=True)

# 5.1 Income Tax
add_calculator_section(
    '5.1 Income Tax Calculator',
    'merasip.com/calculators/income-tax', 'FY 2025-26',
    'Compares Old vs New tax regime with current FY 2025-26 tax slabs, and recommends which regime saves more tax.',
    ['Tax planning season (January-March)',
     'New financial year discussions',
     '"Old regime or new regime?"'],
    [['Gross Annual Salary', 'Rs 3L - 1 Cr'],
     ['HRA Received & Rent Paid', 'Annual amounts'],
     ['City Type', 'Metro (50% HRA) / Non-Metro (40% HRA)'],
     ['Section 80C', 'Rs 0 - 1.5L'],
     ['Section 80D', 'Rs 0 - 1L'],
     ['Home Loan Interest 24(b)', 'Rs 0 - 5L'],
     ['NPS 80CCD(1B)', 'Rs 0 - 50K'],
     ['Other Deductions', 'Rs 0 - 2L']],
    ['Recommended Regime with exact savings amount',
     'Old vs New Regime tax breakdown (slab-wise)',
     'Best Take-Home Salary',
     'Bar chart & Pie charts: regime comparison',
     'Deduction utilization percentages',
     'Dynamic Tax Savings Insights'],
    '"At Rs 12 lakh income with Rs 1.5 lakh in 80C and Rs 25K in 80D, Old Regime saves you Rs 18,200 more. But you have Rs 70,000 unused in 80C. Let me show you how ELSS mutual funds give you both tax saving AND wealth creation..."'
)

# 5.2 Capital Gains Tax
add_calculator_section(
    '5.2 Capital Gains Tax Calculator',
    'merasip.com/calculators/capital-gains-tax', 'Budget 2024',
    'Calculates STCG/LTCG tax on equity, debt, property, and gold with post-Budget 2024 rules.',
    ['Client selling investments and wants to know tax impact',
     'Budget changes explanation',
     'Portfolio rebalancing discussions'],
    [['Asset Type', 'Equity Shares / Equity MF / Debt MF / Property / Gold (tabs)'],
     ['Purchase & Sale Price', 'Rs 10K - 10 Cr each'],
     ['Purchase & Sale Year', 'Year dropdowns'],
     ['Tax Slab Rate', 'For debt/property/gold']],
    ['LTCG or STCG classification with applicable tax rate',
     'Taxable Gain (after Rs 1.25L exemption for equity)',
     'Total Tax Payable (including surcharge and cess)',
     'Cross-asset comparison table',
     'For Property: 20% with indexation vs 12.5% without comparison'],
    '"If you sell your equity mutual fund after 1 year, gains up to Rs 1.25 lakhs are completely tax-free. Beyond that, it\'s just 12.5% -- much lower than FD interest which is taxed at your slab rate of 30%. This is why equity mutual funds are more tax-efficient."'
)

# 5.3 Tax Saving
add_calculator_section(
    '5.3 Tax Saving Calculator',
    'merasip.com/calculators/tax-saving', 'Save Tax',
    'Plans deductions under 80C, 80D, 80CCD(1B), and 24(b) to minimize total tax liability. Shows utilized vs remaining limits.',
    ['Annual tax planning',
     '"How can I save more tax?"',
     'Recommending ELSS, NPS, or health insurance'],
    [['Tax Regime', 'Old / New toggle'],
     ['Gross Income', 'Rs 3L - 1 Cr'],
     ['Section 80C (8 sub-items)', 'EPF, PPF, ELSS, Insurance, Home Loan, Tuition, NSC, Sukanya'],
     ['Section 80D', 'Self & Parents premiums, Senior Citizen toggles'],
     ['NPS 80CCD(1B)', 'Rs 0 - 50K'],
     ['Section 24(b)', 'Rs 0 - 2L']],
    ['Total Tax Saved and Total Deductions',
     'Remaining Limits (unused potential)',
     'Utilization progress bars per section (% used with color coding)',
     'Old vs New regime comparison chart',
     'Smart Recommendations (dynamic tips based on gaps)'],
    '"You\'ve used Rs 80,000 of your Rs 1.5 lakh 80C limit. You still have Rs 70,000 to invest in ELSS. That will save you Rs 21,000 in tax AND earn you 12-15% returns over time. Plus, you haven\'t touched your Rs 50,000 NPS deduction at all."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 6: INSURANCE CALCULATORS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('6. Category D: Insurance Calculators', level=1)
add_body('Navigate to: www.merasip.com/calculators#insurance', color=RGBColor(0x0B, 0x6E, 0xBF), size=11)
add_body('3 Calculators | Help clients understand their insurance gaps and recommend adequate coverage.', bold=True)

# 6.1 HLV
add_calculator_section(
    '6.1 Human Life Value (HLV) Calculator',
    'merasip.com/calculators/human-life-value', 'Life Cover',
    'Calculates the present value of all future income minus personal expenses -- this is the ideal life insurance cover using the Income Replacement Method.',
    ['Life insurance needs assessment',
     '"How much life cover do I need?"',
     'Replacing LIC endowment with term insurance discussions'],
    [['Current Age & Retirement Age', '18-65 / 45-70'],
     ['Annual Income & Growth Rate', 'Rs 1L - 2 Cr / 0-15%'],
     ['Annual Personal Expenses', 'Rs 50K - 1 Cr'],
     ['Loans', 'Home, Car, Other outstanding amounts'],
     ['Children Education & Marriage', 'Future goal amounts'],
     ['Existing Life Cover & Savings', 'Current insurance and investments'],
     ['Discount Rate', '5-12%']],
    ['Recommended Life Insurance Cover (HLV)',
     'Coverage Gap or "Adequately Covered"',
     'Income Multiple (e.g., "15x your annual income")',
     'Coverage Adequacy gauge (progress bar)',
     'HLV component breakdown chart',
     'Year-by-year income projection table'],
    '"Your Human Life Value is Rs 2.3 crores. You currently have Rs 50 lakhs of cover. That\'s a gap of Rs 1.8 crores. If something happens to you, your family\'s lifestyle, children\'s education, and home loan are all at risk. A term plan of Rs 2 crore costs only Rs 15,000/year at your age."'
)

# 6.2 Term Insurance
add_calculator_section(
    '6.2 Term Insurance Calculator',
    'merasip.com/calculators/term-insurance', 'Protection',
    'Uses a Needs-Based approach to calculate adequate term insurance cover -- considering family expenses, goals, loans, and existing resources.',
    ['Term insurance recommendation',
     'Clients who think "Rs 50 lakhs cover is enough"',
     'Annual insurance review'],
    [['Monthly Family Expenses', 'Rs 10K - 5L'],
     ['Number of Dependents', '0-10 (button selector)'],
     ['Years to Support Family', '5-40'],
     ['Inflation & Return Rate', '3-10% / 6-12%'],
     ['Children Education & Marriage', 'Rs 0 - 1 Cr / Rs 0 - 50L'],
     ['Outstanding Loans', 'Home & Other loans'],
     ['Existing Insurance & Savings', 'Current cover & investments'],
     ['Spouse\'s Annual Income', 'Rs 0 - 50L']],
    ['Recommended Sum Assured',
     'Coverage Status: Adequate / Needs Review / Critically Under-Covered',
     'Needs breakdown: Family Expenses + Education + Marriage + Loans',
     'Resources offset: Insurance + Savings + Spouse Income',
     'Coverage Gap with progress bar',
     'Expert Tips section'],
    '"Your family needs Rs 1.8 crore of protection. After accounting for your Rs 40L savings and spouse\'s income, the gap is Rs 1.1 crore. A pure term plan for this amount costs around Rs 12,000/year -- that\'s Rs 33/day to ensure your family is fully protected."'
)

# 6.3 Health Insurance
add_calculator_section(
    '6.3 Health Insurance Calculator',
    'merasip.com/calculators/health-insurance', 'Health Cover',
    'Calculates recommended health insurance cover based on city tier, family size, age, hospital preference, and medical inflation.',
    ['Health insurance planning',
     '"Is my Rs 5 lakh cover enough?"',
     'Employer-only cover review'],
    [['City Tier', 'Tier 1 (Metro) / Tier 2 / Tier 3'],
     ['Family Size', 'Self / Self+Spouse / Family 2+1 / Family 2+2 / Parents 60+'],
     ['Hospital Preference', 'Economy / Standard / Premium / Super Premium'],
     ['Pre-existing Conditions', 'Yes / No'],
     ['Eldest Member Age', '18-80'],
     ['Medical Inflation', '8-18%'],
     ['Planning Horizon', '5-30 years'],
     ['Existing Health Cover', 'Rs 0 - 1 Cr']],
    ['Recommended Coverage Today and Coverage Gap',
     'Future coverage needs (5, 10, 15, 20 years ahead)',
     'City Tier comparison chart',
     'Medical inflation projection chart',
     'Coverage composition pie chart',
     '4 recommendation cards'],
    '"In Mumbai, a family of 4 needs at least Rs 20 lakhs of health cover today. With 14% medical inflation, you\'ll need Rs 55 lakhs in 10 years. Your employer\'s Rs 5 lakh cover is barely enough for one hospital bill. A Rs 15 lakh base plan + Rs 25 lakh super top-up costs only Rs 18,000/year."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 7: LIFE DECISION CALCULATORS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('7. Category E: Life Decision Calculators', level=1)
add_body('Navigate to: www.merasip.com/calculators#life-decision', color=RGBColor(0x0B, 0x6E, 0xBF), size=11)
add_body('5 Calculators | Help clients make better financial decisions in everyday life situations.', bold=True)

# 7.1 Emergency Fund
add_calculator_section(
    '7.1 Emergency Fund Calculator',
    'merasip.com/calculators/emergency-fund', 'Safety Net',
    'Calculates the ideal emergency fund based on monthly expenses, dependents, and job stability.',
    ['Financial planning foundation (this should be step 1)',
     '"How much liquid cash should I keep?"',
     'Before starting SIP -- ensuring basics are covered'],
    [['Monthly Expenses', 'Rent/EMI, Groceries, Utilities, Insurance, Other'],
     ['Number of Dependents', '0-5 (button selector)'],
     ['Job Stability', 'Govt/Stable (3 mo), Corporate (6 mo), Startup (9 mo), Freelance (12 mo)'],
     ['Existing Emergency Savings', 'Rs 0 - 1 Cr'],
     ['Monthly Savings Capacity', 'Rs 0 - 5L']],
    ['Recommended Fund (months x expenses)',
     'Gap to Fill and Months to Build',
     'Adequacy gauge (Fully Funded / Partially Funded / Under-Funded)',
     'Expense breakdown pie chart',
     'Savings timeline table'],
    '"Before we start any SIP, let\'s make sure your foundation is solid. With Rs 45,000 monthly expenses, a corporate job, and 2 dependents, you need Rs 3.6 lakhs as emergency fund. Let\'s build the remaining gap in a liquid fund first, then start your wealth-building SIP."'
)

# 7.2 Net Worth
add_calculator_section(
    '7.2 Net Worth Calculator',
    'merasip.com/calculators/net-worth', 'Wealth Score',
    'Tracks all your assets and liabilities to calculate total net worth and gives a Wealth Score from 0-100.',
    ['Annual financial health check',
     '"Am I doing well financially?"',
     'Comprehensive financial review'],
    [['Assets (add/remove)', 'Savings, FDs, Mutual Funds, Stocks, PPF, Real Estate, Gold, Vehicle etc. Mark each as "Liquid" or not'],
     ['Liabilities (add/remove)', 'Home Loan, Car Loan, Credit Card, Personal Loan etc.']],
    ['Net Worth (total assets minus total liabilities)',
     'Liquid Net Worth (liquid assets minus liabilities)',
     'Debt-to-Asset Ratio',
     'Wealth Score (0-100): Excellent / Good / Fair / Poor / Critical',
     'Assets vs Liabilities bar chart',
     'Asset and Liability breakdown pie charts'],
    '"Your net worth is Rs 38 lakhs, but Rs 30 lakhs is your flat -- that\'s illiquid. Your liquid net worth is only Rs 8 lakhs against Rs 15 lakhs in loans. Your wealth score is 62/100 -- Fair. Let\'s work on improving this by building liquid assets through SIPs."'
)

# 7.3 Rent vs Buy
add_calculator_section(
    '7.3 Rent vs Buy Calculator',
    'merasip.com/calculators/rent-vs-buy', 'Decision',
    'Compares the total financial impact of renting (and investing the savings) vs buying a home over time, with break-even analysis.',
    ['Client deciding whether to buy or continue renting',
     '"Is buying a house a good investment?"',
     'Young professionals in expensive cities'],
    [['Property Price', 'Rs 10L - 50 Cr'],
     ['Down Payment & Loan Rate', '0-100% / 5-15%'],
     ['Loan Tenure', '5-30 years'],
     ['Monthly Rent & Annual Increase', 'Rs 1K-5L / 0-15%'],
     ['Property Appreciation', '0-15%'],
     ['Maintenance & Property Tax', '0-5% / 0-3%'],
     ['Investment Return Rate', '5-20%'],
     ['Time Horizon', '5-30 years']],
    ['Verdict: "Buy" or "Continue Renting & Invest"',
     'Savings amount and Break-even year',
     'Cumulative cost comparison chart',
     'Year-by-year breakdown table'],
    '"For a Rs 80 lakh flat in Bangalore at 8.5% loan vs Rs 25,000 rent, buying breaks even in year 11. If you plan to stay 15+ years, buying wins. But if you might relocate in 5-7 years, renting and investing the difference actually gives you Rs 12 lakhs more."'
)

# 7.4 FD vs Loan
add_calculator_section(
    '7.4 Break FD vs Take Loan Calculator',
    'merasip.com/calculators/fd-vs-loan', 'Decision',
    'When you need cash urgently, should you break your Fixed Deposit or take a loan? Calculates the net cost of both options including tax impact.',
    ['Client needs emergency funds',
     '"Should I break my FD?"',
     'Liquidity planning discussions'],
    [['FD Amount & Interest Rate', 'Rs 10K-10Cr / 3-10%'],
     ['FD Remaining Tenure', '0.5-10 years'],
     ['Premature Withdrawal Penalty', '0-3%'],
     ['Amount Needed', 'Rs 10K - 10 Cr'],
     ['Loan Interest Rate & Tenure', '5-25% / 0.5-10 years'],
     ['Processing Fee', '0-5%'],
     ['Tax Bracket', '0% / 5% / 10% / 15% / 20% / 30%']],
    ['Verdict: "Break the FD" or "Take the Loan"',
     'Savings amount',
     'Cost comparison bar chart',
     'Detailed 8-metric comparison table'],
    '"Your FD of Rs 5 lakhs at 7% costs Rs 18,500 to break. A personal loan at 12% costs Rs 33,200 in interest + fees. Breaking the FD saves Rs 19,700. But consider this -- if that money were in a liquid mutual fund, you\'d have instant access with no penalty."'
)

# 7.5 Lifestyle Inflation
add_calculator_section(
    '7.5 Lifestyle Inflation Calculator',
    'merasip.com/calculators/lifestyle-inflation', 'Eye-Opener',
    'Shows how "lifestyle creep" -- increasing spending faster than income -- gradually destroys your savings rate over time.',
    ['Clients who earn well but save poorly',
     '"Where does my money go?"',
     'Behavioral finance discussions'],
    [['Monthly Income & Expenses', 'Current amounts'],
     ['Annual Income Growth', '0-30%'],
     ['Annual Lifestyle Inflation', '0-30%'],
     ['Projection Years', '5-40'],
     ['Scenario', 'Your Plan / Disciplined (6%) / Aggressive Saver (0%)']],
    ['Current and Final Year Savings Rate',
     'Crossover Year (when expenses overtake income -- danger!)',
     '3-scenario comparison chart',
     'Income vs Expenses area chart',
     'Year-by-year table with colored savings rate badges'],
    '"You earn Rs 1 lakh and spend Rs 65,000 -- a healthy 35% savings rate. But if expenses grow at 12% while income grows at 10%, by year 8 your savings rate drops to 15%, and by year 14 your expenses overtake income! Keep lifestyle inflation to 6% and you\'ll have Rs 45 lakhs more."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 8: CALCULATOR-TO-PRODUCT MAPPING
# ═══════════════════════════════════════════════════════════════

doc.add_heading('8. Calculator-to-Product Mapping', level=1)
add_body('Use this quick reference to connect calculators to product recommendations:', bold=True)

add_table(
    ['Calculator', 'Product to Recommend'],
    [
        ['SIP, Step-Up SIP, Goal-Based', 'Equity Mutual Fund SIP'],
        ['Inflation-Adjusted, Cost of Delay', 'Start SIP TODAY (urgency)'],
        ['Retirement, FIRE', 'Long-term Equity + Balanced Funds'],
        ['SWP', 'Balanced Advantage Fund for withdrawal'],
        ['Lumpsum', 'Equity or Debt Fund (based on horizon)'],
        ['EMI, Loan Prepayment', 'Advisory only (builds trust)'],
        ['Car Loan vs Cash, Rent vs Buy', 'Lumpsum + SIP combo recommendation'],
        ['Home Affordability', 'Advisory + Home Loan referral'],
        ['Income Tax, Tax Saving', 'ELSS Mutual Funds + NPS'],
        ['Capital Gains', 'Tax Harvesting + Portfolio Review'],
        ['HLV, Term Insurance', 'Term Insurance plan'],
        ['Health Insurance', 'Health Insurance + Super Top-Up'],
        ['Emergency Fund', 'Liquid Fund / Money Market Fund'],
        ['Net Worth', 'Comprehensive portfolio review'],
        ['FD vs Loan', 'Liquid Fund over FD recommendation'],
        ['Lifestyle Inflation', 'Step-Up SIP + Disciplined spending'],
        ['Life-Stage, Lifeline', 'Complete Financial Plan'],
        ['Market Correction', 'Continue SIP (don\'t panic)'],
    ],
    col_widths=[3.0, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 9: CLIENT CONVERSATION SCRIPTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('9. Client Conversation Scripts', level=1)
add_body('Ready-to-use conversation flows for different meeting types:', bold=True)

doc.add_heading('Script 1: First Meeting -- The "Eye-Opener" Approach', level=2)
steps = [
    'Start with Cost of Delay Calculator -- create urgency',
    'Move to SIP Calculator -- show the opportunity',
    'Use Goal-Based Calculator -- connect to their life goals',
    'Close with Step-Up SIP -- show how small increases make huge differences',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'Step {i}: {step}')

doc.add_heading('Script 2: Tax Season -- The "Save Tax & Build Wealth" Approach', level=2)
steps = [
    'Start with Income Tax Calculator -- find their regime and gap',
    'Move to Tax Saving Calculator -- identify unused limits',
    'Recommend ELSS SIP to fill 80C gap',
    'Show SIP Calculator -- project their ELSS returns beyond the lock-in',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'Step {i}: {step}')

doc.add_heading('Script 3: Life Insurance Review', level=2)
steps = [
    'Start with Human Life Value Calculator -- show the number',
    'Move to Term Insurance Calculator -- needs-based calculation',
    'Identify coverage gap',
    'Quote a term plan to fill the gap',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'Step {i}: {step}')

doc.add_heading('Script 4: Home Buyer Consultation', level=2)
steps = [
    'Start with Home Affordability Calculator -- set realistic budget',
    'Use EMI Calculator -- show exact monthly outflow',
    'Show Rent vs Buy Calculator -- validate the decision',
    'Add Loan Prepayment Calculator -- show how to save on interest',
    'Suggest SIP alongside home loan for wealth creation',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'Step {i}: {step}')

doc.add_heading('Script 5: Comprehensive Financial Review', level=2)
steps = [
    'Net Worth Calculator -- current position',
    'Emergency Fund Calculator -- safety net check',
    'Term Insurance Calculator -- protection check',
    'Health Insurance Calculator -- health cover check',
    'Retirement Planner -- long-term plan',
    'SIP Calculator -- start the investment journey',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'Step {i}: {step}')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 10: COMMON CLIENT SCENARIOS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('10. Common Client Scenarios & Objection Handling', level=1)

scenarios = [
    ('"I don\'t have money to invest"',
     'Use Daily SIP Calculator: Show Rs 100/day = Rs 14 lakhs in 10 years. Use Lifestyle Inflation Calculator: Show where spending leaks are.'),
    ('"Markets are too risky / Markets are crashing"',
     'Use Market Correction Calculator: Show that even a 30% crash has only 8% long-term impact. Use SIP vs Lumpsum: Show rupee cost averaging benefit.'),
    ('"I\'ll start next year"',
     'Use Cost of Delay Calculator: Show Rs 1,000/day of delay cost. This is your most powerful closer.'),
    ('"FD is safer"',
     'Use SIP Calculator at 12% vs Inflation-Adjusted at 7% inflation. Show that FD barely beats inflation. Use Capital Gains Calculator: Show FD interest taxed at 30% vs equity LTCG at 12.5%.'),
    ('"I already have LIC"',
     'Use HLV Calculator: Show total cover needed. Use Term Insurance Calculator: Show the gap. Explain that LIC endowment gives 5% return vs SIP at 12%.'),
    ('"My employer gives health insurance"',
     'Use Health Insurance Calculator: Show that Rs 3-5L employer cover is insufficient. Show medical inflation projection -- Rs 5L surgery today = Rs 14L in 10 years.'),
    ('"I want to retire early"',
     'Use FIRE Calculator: Calculate their FIRE Number. Use Life-Stage Planner: Map the complete journey. Use Step-Up SIP: Show how incremental discipline accelerates FIRE.'),
]

for objection, response in scenarios:
    doc.add_heading(objection, level=3)
    add_body(response)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 11: TIPS FOR POSPs
# ═══════════════════════════════════════════════════════════════

doc.add_heading('11. Tips for POSPs & Sub-Brokers', level=1)

doc.add_heading("Do's", level=2)
dos = [
    'Always use the client\'s real numbers -- never use generic examples when you have their actual data',
    'Download PDFs -- give the client a tangible takeaway from every meeting',
    'Start with a problem -- use Insurance, Tax, or Emergency Fund calculators to identify gaps, then present SIP as the solution',
    'Show multiple scenarios -- "What if you invest Rs 10K? Rs 15K? Rs 20K?"',
    'Use Cost of Delay in every meeting -- it creates natural urgency without being pushy',
    'Bookmark your top 5 calculators on your phone for instant access during meetings',
    'Follow the conversation flow: Problem (gap identification) --> Solution (SIP/Insurance) --> Action (start today)',
]
for item in dos:
    add_bullet(item)

doc.add_heading("Don'ts", level=2)
donts = [
    'Don\'t promise exact returns -- always say "assuming X% returns" and point to the disclaimer',
    'Don\'t skip the disclaimer -- every calculator shows it; respect it',
    'Don\'t overwhelm -- use 2-3 calculators per meeting, not all 30',
    'Don\'t use calculators as the ONLY sales tool -- they support the conversation, not replace it',
    'Don\'t misrepresent -- these are educational tools, not guarantees',
]
for item in donts:
    add_bullet(item)

doc.add_heading('Recommended Calculator Combos per Meeting Type', level=2)
add_table(
    ['Meeting Type', 'Primary Calculator', 'Supporting Calculators'],
    [
        ['First Meeting', 'Cost of Delay', 'SIP, Goal-Based'],
        ['Tax Planning', 'Income Tax', 'Tax Saving, Capital Gains'],
        ['Insurance Review', 'HLV or Term Insurance', 'Health Insurance'],
        ['Home Buying', 'Home Affordability', 'EMI, Rent vs Buy'],
        ['Retirement Planning', 'Retirement Planner', 'FIRE, Life-Stage'],
        ['Annual Review', 'Net Worth', 'Lifestyle Inflation, Step-Up SIP'],
        ['Market Panic', 'Correction Impact', 'SIP (continue investing)'],
    ],
    col_widths=[1.8, 2.0, 2.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION 12: QUICK REFERENCE URLs
# ═══════════════════════════════════════════════════════════════

doc.add_heading('12. Quick Reference: All 30 Calculator URLs', level=1)

doc.add_heading('Wealth & SIP (15)', level=2)
add_table(
    ['#', 'Calculator', 'URL'],
    [
        ['1', 'SIP Future Value', 'merasip.com/calculators/sip'],
        ['2', 'Step-Up SIP', 'merasip.com/calculators/step-up-sip'],
        ['3', 'Goal-Based SIP', 'merasip.com/calculators/goal-based'],
        ['4', 'Inflation-Adjusted SIP', 'merasip.com/calculators/inflation-adjusted'],
        ['5', 'Retirement SIP Planner', 'merasip.com/calculators/retirement'],
        ['6', 'SWP Calculator', 'merasip.com/calculators/swp'],
        ['7', 'Lumpsum Calculator', 'merasip.com/calculators/lumpsum'],
        ['8', 'SIP vs Lumpsum', 'merasip.com/calculators/sip-vs-lumpsum'],
        ['9', 'Duration Optimizer', 'merasip.com/calculators/duration-optimizer'],
        ['10', 'Market Correction Impact', 'merasip.com/calculators/correction-impact'],
        ['11', 'Life-Stage Planner', 'merasip.com/calculators/life-stage'],
        ['12', 'Lifeline Financial Planner', 'merasip.com/calculators/lifeline'],
        ['13', 'Daily SIP', 'merasip.com/calculators/daily-sip'],
        ['14', 'FIRE Calculator', 'merasip.com/calculators/fire'],
        ['15', 'Cost of Delay', 'merasip.com/calculators/delay-cost'],
    ],
    col_widths=[0.4, 2.3, 3.5]
)

doc.add_heading('Loan (4)', level=2)
add_table(
    ['#', 'Calculator', 'URL'],
    [
        ['16', 'EMI Calculator', 'merasip.com/calculators/emi'],
        ['17', 'Loan Prepayment', 'merasip.com/calculators/loan-prepayment'],
        ['18', 'Car Loan vs Cash', 'merasip.com/calculators/car-loan-vs-cash'],
        ['19', 'Home Affordability', 'merasip.com/calculators/home-affordability'],
    ],
    col_widths=[0.4, 2.3, 3.5]
)

doc.add_heading('Tax (3)', level=2)
add_table(
    ['#', 'Calculator', 'URL'],
    [
        ['20', 'Income Tax (Old vs New)', 'merasip.com/calculators/income-tax'],
        ['21', 'Capital Gains Tax', 'merasip.com/calculators/capital-gains-tax'],
        ['22', 'Tax Saving (80C/80D)', 'merasip.com/calculators/tax-saving'],
    ],
    col_widths=[0.4, 2.3, 3.5]
)

doc.add_heading('Insurance (3)', level=2)
add_table(
    ['#', 'Calculator', 'URL'],
    [
        ['23', 'Human Life Value', 'merasip.com/calculators/human-life-value'],
        ['24', 'Term Insurance', 'merasip.com/calculators/term-insurance'],
        ['25', 'Health Insurance', 'merasip.com/calculators/health-insurance'],
    ],
    col_widths=[0.4, 2.3, 3.5]
)

doc.add_heading('Life Decisions (5)', level=2)
add_table(
    ['#', 'Calculator', 'URL'],
    [
        ['26', 'Emergency Fund', 'merasip.com/calculators/emergency-fund'],
        ['27', 'Net Worth', 'merasip.com/calculators/net-worth'],
        ['28', 'Rent vs Buy', 'merasip.com/calculators/rent-vs-buy'],
        ['29', 'Break FD vs Take Loan', 'merasip.com/calculators/fd-vs-loan'],
        ['30', 'Lifestyle Inflation', 'merasip.com/calculators/lifestyle-inflation'],
    ],
    col_widths=[0.4, 2.3, 3.5]
)

# ═══════════════════════════════════════════════════════════════
# REGULATORY NOTE & FOOTER
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
add_section_divider()

p = doc.add_paragraph()
run = p.add_run('Important Regulatory Note:')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_body('All calculators are for educational purposes only. They are not investment advice. Past performance does not guarantee future results. Mutual fund investments are subject to market risks. Read all scheme-related documents carefully. Insurance is subject to IRDAI regulations. Tax calculations are based on current tax laws and may change.', size=9, color=RGBColor(0x66, 0x66, 0x66))

add_section_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document Version: 1.0  |  Last Updated: March 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for internal use by Trustner Asset Services Pvt. Ltd.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_path = '/Users/ram/Documents/Trustner Tech Project/Mera SIPOnline/MeraSIP-Calculator-Training-Guide.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
