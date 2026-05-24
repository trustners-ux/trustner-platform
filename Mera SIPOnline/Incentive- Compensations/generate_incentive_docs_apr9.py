#!/usr/bin/env python3
"""
Generate Trustner Incentive System Reference Documents — Dated 9 April 2026
Outputs:
  1. Trustner_Incentive_Reference_9Apr2026.xlsx  (Excel workbook, multiple sheets)
  2. Trustner_Incentive_Reference_9Apr2026.docx  (Word document)
"""

import os
from datetime import date

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_DATE = "9 April 2026"
DOC_DATE_SHORT = "9Apr2026"

# ─── Colour palette ───
NAVY = "1B2A4A"
GOLD = "D4A843"
WHITE = "FFFFFF"
LIGHT_GREY = "F2F2F2"
GREEN = "27AE60"
RED = "C0392B"

# ─── Style helpers — Excel ───
header_font = Font(name="Calibri", bold=True, size=11, color=WHITE)
header_fill = PatternFill(start_color=NAVY, end_color=NAVY, fill_type="solid")
gold_fill = PatternFill(start_color=GOLD, end_color=GOLD, fill_type="solid")
alt_fill = PatternFill(start_color=LIGHT_GREY, end_color=LIGHT_GREY, fill_type="solid")
title_font = Font(name="Calibri", bold=True, size=14, color=NAVY)
section_font = Font(name="Calibri", bold=True, size=12, color=NAVY)
thin_border = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)

def style_header_row(ws, row, cols):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

def style_data_row(ws, row, cols, alt=False):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        if alt:
            cell.fill = alt_fill

def auto_width(ws):
    for col_cells in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col_cells[0].column)
        for cell in col_cells:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)


# ════════════════════════════════════════════════════════════════════
#  EXCEL WORKBOOK
# ════════════════════════════════════════════════════════════════════
def build_excel():
    wb = openpyxl.Workbook()

    # ── Sheet 1: Cover ──
    ws = wb.active
    ws.title = "Cover"
    ws.merge_cells("A2:F2")
    ws["A2"] = "TRUSTNER ASSET SERVICES PVT LTD"
    ws["A2"].font = Font(name="Calibri", bold=True, size=18, color=NAVY)
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A4:F4")
    ws["A4"] = "INCENTIVE & COMPENSATION FRAMEWORK"
    ws["A4"].font = Font(name="Calibri", bold=True, size=16, color=GOLD)
    ws["A4"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A6:F6")
    ws["A6"] = f"Reference Document — {DOC_DATE}"
    ws["A6"].font = Font(name="Calibri", size=12, color=NAVY)
    ws["A6"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A8:F8")
    ws["A8"] = "ARN-286886 | IRDAI Corporate Agent (Composite) — Applied"
    ws["A8"].font = Font(name="Calibri", size=11, color="666666")
    ws["A8"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A10:F10")
    ws["A10"] = "CONTENTS"
    ws["A10"].font = section_font
    contents = [
        "1. DST (Direct Sales Team) Incentive Slabs",
        "2. POSP RM / RM POSP / RM Wealth Incentive Slabs",
        "3. Product Tier & Weightage System",
        "4. 10-Step Incentive Calculation Formula",
        "5. Channel Payout Structure",
        "6. POSP Category System (4 LOBs × 9 Categories)",
        "7. Motor Insurance Commission Framework",
        "8. Trail Income Distribution",
        "9. SIP & Lumpsum Clawback Rules",
        "10. MF Data Approval Flow",
        "11. 6-Step Payout Approval Workflow",
        "12. POSP Coded Naming Convention",
        "13. Area Manager / Regional Head (Deferred to Q2 FY27)",
        "14. Key Removals for Q1 FY27",
        "15. Partner Hierarchy & Channel Types",
        "16. Franchise Differential Payout Model",
        "17. Referral Code System",
    ]
    for i, c in enumerate(contents):
        ws[f"A{12+i}"] = c
        ws[f"A{12+i}"].font = Font(name="Calibri", size=11)

    # ── Sheet 2: DST Slabs ──
    ws2 = wb.create_sheet("DST Slabs")
    ws2.merge_cells("A1:E1")
    ws2["A1"] = "DST (Direct Sales Team) Incentive Slabs"
    ws2["A1"].font = title_font
    ws2.merge_cells("A2:E2")
    ws2["A2"] = "No multiplier concept — flat slab-based incentive on weighted business"
    ws2["A2"].font = Font(name="Calibri", italic=True, size=10, color="666666")

    headers = ["Slab", "Achievement Range", "Incentive %", "Tier Name", "Description"]
    for i, h in enumerate(headers, 1):
        ws2.cell(row=4, column=i, value=h)
    style_header_row(ws2, 4, 5)

    data = [
        [1, "0 – 80%", "0%", "No Incentive", "Below threshold — no payout"],
        [2, "80.01 – 100%", "4%", "Base", "Meets minimum threshold"],
        [3, "100.01 – 125%", "5%", "Enhanced", "Exceeds target"],
        [4, "125.01 – 150%", "6%", "Super", "Significantly above target"],
        [5, "150.01% +", "8%", "Champion", "Outstanding performance"],
    ]
    for r, row_data in enumerate(data, 5):
        for c, val in enumerate(row_data, 1):
            ws2.cell(row=r, column=c, value=val)
        style_data_row(ws2, r, 5, alt=(r % 2 == 0))

    ws2.merge_cells("A11:E11")
    ws2["A11"] = "Formula: Incentive = Weighted Business × Slab %"
    ws2["A11"].font = Font(name="Calibri", bold=True, size=11, color=GREEN)

    ws2.merge_cells("A13:E13")
    ws2["A13"] = "Example: DST achieves 112% → Slab = 5% (Enhanced)"
    ws2["A13"].font = Font(name="Calibri", size=10)
    ws2.merge_cells("A14:E14")
    ws2["A14"] = "If weighted business = ₹15,00,000 → Incentive = ₹15,00,000 × 5% = ₹75,000"
    ws2["A14"].font = Font(name="Calibri", size=10)
    auto_width(ws2)

    # ── Sheet 3: POSP RM Slabs ──
    ws3 = wb.create_sheet("POSP RM Slabs")
    ws3.merge_cells("A1:E1")
    ws3["A1"] = "POSP RM / RM POSP / RM Wealth Incentive Slabs"
    ws3["A1"].font = title_font
    ws3.merge_cells("A2:E2")
    ws3["A2"] = "Applicable to all RM roles managing POSP/sub-broker channels — no multiplier"
    ws3["A2"].font = Font(name="Calibri", italic=True, size=10, color="666666")

    for i, h in enumerate(headers, 1):
        ws3.cell(row=4, column=i, value=h)
    style_header_row(ws3, 4, 5)

    data3 = [
        [1, "0 – 80%", "0%", "No Incentive", "Below threshold — no payout"],
        [2, "80.01 – 100%", "1.2%", "Base", "Meets minimum threshold"],
        [3, "100.01 – 125%", "1.5%", "Enhanced", "Exceeds target"],
        [4, "125.01 – 150%", "1.8%", "Super", "Significantly above target"],
        [5, "150.01% +", "2.4%", "Champion", "Outstanding performance"],
    ]
    for r, row_data in enumerate(data3, 5):
        for c, val in enumerate(row_data, 1):
            ws3.cell(row=r, column=c, value=val)
        style_data_row(ws3, r, 5, alt=(r % 2 == 0))

    ws3.merge_cells("A11:E11")
    ws3["A11"] = "Formula: Incentive = Weighted Business × Slab %"
    ws3["A11"].font = Font(name="Calibri", bold=True, size=11, color=GREEN)

    ws3.merge_cells("A13:E13")
    ws3["A13"] = "Note: POSP RM manages POSP agents; their incentive is on the team's weighted business."
    ws3["A13"].font = Font(name="Calibri", size=10)
    auto_width(ws3)

    # ── Sheet 4: Product Tiers ──
    ws4 = wb.create_sheet("Product Tiers")
    ws4.merge_cells("A1:E1")
    ws4["A1"] = "Product Tier & Weightage System"
    ws4["A1"].font = title_font
    ws4.merge_cells("A2:E2")
    ws4["A2"] = 'Weightage determines how much of the premium counts toward target achievement (renamed from "Credit")'
    ws4["A2"].font = Font(name="Calibri", italic=True, size=10, color="666666")

    tier_headers = ["Tier", "Commission Earned", "Weightage %", "Example Products", "Weighted Business Formula"]
    for i, h in enumerate(tier_headers, 1):
        ws4.cell(row=4, column=i, value=h)
    style_header_row(ws4, 4, 5)

    tier_data = [
        ["Tier 1", "≥ 30%", "100%", "Traditional Life, Health (High Commission)", "Premium × 100%"],
        ["Tier 2", "15% – 30%", "75%", "ULIP, Standard Health, Fire", "Premium × 75%"],
        ["Tier 3", "< 15%", "50%", "Term Life, Motor, Standard GI", "Premium × 50%"],
        ["Tier 4", "≤ 7.5%", "25%", "Group Term, Group Mediclaim, Accidental", "Premium × 25%"],
    ]
    for r, row_data in enumerate(tier_data, 5):
        for c, val in enumerate(row_data, 1):
            ws4.cell(row=r, column=c, value=val)
        style_data_row(ws4, r, 5, alt=(r % 2 == 0))

    ws4.merge_cells("A10:E10")
    ws4["A10"] = "Weighted Business = Σ (Each Policy Premium × Tier Weightage %)"
    ws4["A10"].font = Font(name="Calibri", bold=True, size=11, color=GREEN)

    ws4.merge_cells("A12:E12")
    ws4["A12"] = "Example: ₹1,00,000 Tier 1 + ₹2,00,000 Tier 3 = ₹1,00,000 + ₹1,00,000 = ₹2,00,000 weighted"
    ws4["A12"].font = Font(name="Calibri", size=10)
    auto_width(ws4)

    # ── Sheet 5: 10-Step Formula ──
    ws5 = wb.create_sheet("Calculation Formula")
    ws5.merge_cells("A1:D1")
    ws5["A1"] = "10-Step Incentive Calculation Pipeline"
    ws5["A1"].font = title_font

    formula_headers = ["Step", "Action", "Formula / Logic", "Output"]
    for i, h in enumerate(formula_headers, 1):
        ws5.cell(row=3, column=i, value=h)
    style_header_row(ws5, 3, 4)

    steps = [
        ["1", "Collect Business Entries", "All policies written in the month", "Raw premium list"],
        ["2", "Assign Product Tier", "Based on commission % earned on product", "Tier 1/2/3/4 per policy"],
        ["3", "Calculate Weighted Business", "Each premium × tier weightage %", "Weighted premium per policy"],
        ["4", "Sum Weighted Business", "Σ all weighted premiums", "Total weighted business (₹)"],
        ["5", "Get Monthly Target", "From employee master / goal sheet", "Target amount (₹)"],
        ["6", "Calculate Achievement %", "(Weighted Business ÷ Target) × 100", "Achievement percentage"],
        ["7", "Look Up Slab", "DST or POSP RM slab table", "Incentive % for the slab"],
        ["8", "Calculate Gross Incentive", "Weighted Business × Slab %", "Gross incentive amount (₹)"],
        ["9", "Apply Deductions (if any)", "Clawbacks, recoveries", "Net incentive amount (₹)"],
        ["10", "Route for Approval", "6-step approval workflow", "Final approved payout"],
    ]
    for r, row_data in enumerate(steps, 4):
        for c, val in enumerate(row_data, 1):
            ws5.cell(row=r, column=c, value=val)
        style_data_row(ws5, r, 4, alt=(r % 2 == 0))
    auto_width(ws5)

    # ── Sheet 6: Channel Payouts ──
    ws6 = wb.create_sheet("Channel Payouts")
    ws6.merge_cells("A1:E1")
    ws6["A1"] = "Channel Payout Structure"
    ws6["A1"].font = title_font

    ch_headers = ["Channel", "Payout to Channel", "Company Margin", "Notes", "Incentive Slab Used"]
    for i, h in enumerate(ch_headers, 1):
        ws6.cell(row=3, column=i, value=h)
    style_header_row(ws6, 3, 5)

    ch_data = [
        ["POSP", "70%", "30%", "Standard POSP agents", "POSP RM Slabs (for managing RM)"],
        ["BQP (Referral)", "85%", "15%", "Business Quality Partner / Referral channel", "—"],
        ["Referral", "50%", "50%", "Casual referral partners", "—"],
        ["Employee (DST)", "0% (salaried)", "100%", "Incentive via DST slabs, not commission share", "DST Slabs"],
        ["Sub-Broker", "60%", "40%", "Registered sub-brokers under Trustner", "—"],
        ["Franchise", "80% – 90%", "10% – 20%", "Range based on agreement & volume", "—"],
    ]
    for r, row_data in enumerate(ch_data, 4):
        for c, val in enumerate(row_data, 1):
            ws6.cell(row=r, column=c, value=val)
        style_data_row(ws6, r, 5, alt=(r % 2 == 0))

    ws6.merge_cells("A11:E11")
    ws6["A11"] = "Note: BQP payout revised from 65% to 85%. Franchise is a range (80-90%), not flat 85%."
    ws6["A11"].font = Font(name="Calibri", italic=True, size=10, color=RED)
    auto_width(ws6)

    # ── Sheet 7: POSP Categories ──
    ws7 = wb.create_sheet("POSP Categories")
    ws7.merge_cells("A1:F1")
    ws7["A1"] = "POSP Category System — 4 LOBs × 9 Categories"
    ws7["A1"].font = title_font
    ws7.merge_cells("A2:F2")
    ws7["A2"] = "Each POSP is assigned a category per LOB independently. Category determines % of company commission passed to POSP."
    ws7["A2"].font = Font(name="Calibri", italic=True, size=10, color="666666")

    posp_headers = ["Category", "Payout % to POSP", "Life", "Health", "GI Non-Motor", "GI Motor"]
    for i, h in enumerate(posp_headers, 1):
        ws7.cell(row=4, column=i, value=h)
    style_header_row(ws7, 4, 6)

    posp_cats = [
        ["A", "50%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["B", "55%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["C", "60%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["D", "65%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["D+", "70%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["E", "75%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["E+", "80%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["F1", "85%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
        ["F2", "90%", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP", "Assigned per POSP"],
    ]
    for r, row_data in enumerate(posp_cats, 5):
        for c, val in enumerate(row_data, 1):
            ws7.cell(row=r, column=c, value=val)
        style_data_row(ws7, r, 6, alt=(r % 2 == 0))

    ws7.merge_cells("A15:F15")
    ws7["A15"] = "4 LOBs: Life Insurance | Health Insurance | GI Non-Motor | GI Motor"
    ws7["A15"].font = Font(name="Calibri", bold=True, size=11, color=NAVY)
    ws7.merge_cells("A16:F16")
    ws7["A16"] = "A POSP can be Category C in Life but Category E in Health — independent per LOB."
    ws7["A16"].font = Font(name="Calibri", size=10)
    ws7.merge_cells("A17:F17")
    ws7["A17"] = "Coded Naming: TIB/POSP/001/2526 (Trustner Insurance Broking / POSP / Serial / FY)"
    ws7["A17"].font = Font(name="Calibri", size=10, color="666666")
    ws7.merge_cells("A19:F19")
    ws7["A19"] = "Category assignment is reviewed manually at end of each quarter by Admin (Sangeeta Shah) / Super Admin (Ram Shah). No automatic progression."
    ws7["A19"].font = Font(name="Calibri", bold=True, size=10, color=RED)
    auto_width(ws7)

    # ── Sheet 8: Motor Framework ──
    ws8 = wb.create_sheet("Motor Insurance")
    ws8.merge_cells("A1:E1")
    ws8["A1"] = "Motor Insurance Commission Framework"
    ws8["A1"].font = title_font

    ws8.merge_cells("A3:E3")
    ws8["A3"] = "OD (Own Damage) / TP (Third Party) Split Calculation"
    ws8["A3"].font = section_font

    motor_headers = ["Component", "Commission Source", "Rate Range", "Who Earns", "Notes"]
    for i, h in enumerate(motor_headers, 1):
        ws8.cell(row=5, column=i, value=h)
    style_header_row(ws8, 5, 5)

    motor_data = [
        ["OD Premium", "Insurer pays on OD portion", "10% – 15%", "Trustner (broker)", "Varies by insurer & vehicle type"],
        ["TP Premium", "Fixed by IRDAI — no commission", "0%", "—", "No commission on TP premium"],
        ["Net Premium", "OD + TP combined", "Blended rate", "—", "Effective rate lower due to TP drag"],
        ["POSP Payout", "From Trustner's OD commission", "Per category %", "POSP", "Category A(50%) to F2(90%) of Trustner's share"],
        ["BQP Payout", "From Trustner's OD commission", "85%", "BQP", "85% of Trustner's commission on OD"],
        ["Franchise Payout", "From Trustner's OD commission", "80-90%", "Franchise", "Range per agreement"],
    ]
    for r, row_data in enumerate(motor_data, 6):
        for c, val in enumerate(row_data, 1):
            ws8.cell(row=r, column=c, value=val)
        style_data_row(ws8, r, 5, alt=(r % 2 == 0))

    ws8.merge_cells("A13:E13")
    ws8["A13"] = "Product Tier: Motor falls under Tier 3 (< 15% commission) → 50% weightage"
    ws8["A13"].font = Font(name="Calibri", bold=True, size=10, color=GREEN)
    ws8.merge_cells("A14:E14")
    ws8["A14"] = "Example: ₹50,000 motor policy (OD ₹30,000 + TP ₹20,000). Commission on OD = 15% = ₹4,500."
    ws8["A14"].font = Font(name="Calibri", size=10)
    ws8.merge_cells("A15:E15")
    ws8["A15"] = "If POSP is Category D (65%): POSP gets ₹4,500 × 65% = ₹2,925. Trustner retains ₹1,575."
    ws8["A15"].font = Font(name="Calibri", size=10)
    auto_width(ws8)

    # ── Sheet 9: Trail Income ──
    ws9 = wb.create_sheet("Trail Income")
    ws9.merge_cells("A1:D1")
    ws9["A1"] = "Trail Income Distribution — Mutual Funds"
    ws9["A1"].font = title_font
    ws9.merge_cells("A2:D2")
    ws9["A2"] = "Management column removed. Split is between RM and Company only."
    ws9["A2"].font = Font(name="Calibri", italic=True, size=10, color=RED)

    trail_headers = ["Source Type", "RM Share", "Company Share", "Condition"]
    for i, h in enumerate(trail_headers, 1):
        ws9.cell(row=4, column=i, value=h)
    style_header_row(ws9, 4, 4)

    trail_data = [
        ["Self-Sourced — New PAN", "25%", "75%", "RM brings a completely new client (new PAN)"],
        ["Self-Sourced — Existing PAN", "15%", "85%", "RM brings new business from existing client"],
        ["Assigned — New Business", "15%", "85%", "Company assigns client, RM generates new business"],
        ["Assigned — No New Business", "0%", "100%", "Company assigns client, no new business generated"],
        ["Office Walk-in", "10%", "90%", "Client walks in, RM services them"],
    ]
    for r, row_data in enumerate(trail_data, 5):
        for c, val in enumerate(row_data, 1):
            ws9.cell(row=r, column=c, value=val)
        style_data_row(ws9, r, 4, alt=(r % 2 == 0))

    ws9.merge_cells("A11:D11")
    ws9["A11"] = "Trail Rates (Annual — from AMC to Trustner ARN-286886):"
    ws9["A11"].font = section_font

    trail_rate_headers = ["Fund Category", "Typical Trail Rate", "Range"]
    for i, h in enumerate(trail_rate_headers, 1):
        ws9.cell(row=13, column=i, value=h)
    style_header_row(ws9, 13, 3)

    trail_rates = [
        ["Equity — Active", "0.70% – 1.50%", "Varies by AMC"],
        ["Equity — Index / Passive", "0.30% – 0.50%", "Lower TER products"],
        ["Hybrid / Balanced", "0.50% – 1.00%", "Mid-range"],
        ["Debt — Active", "0.20% – 0.80%", "Duration dependent"],
        ["Liquid / Overnight", "0.05% – 0.15%", "Very low trail"],
    ]
    for r, row_data in enumerate(trail_rates, 14):
        for c, val in enumerate(row_data, 1):
            ws9.cell(row=r, column=c, value=val)
        style_data_row(ws9, r, 3, alt=(r % 2 == 0))

    ws9.merge_cells("A20:D20")
    ws9["A20"] = "MF data entered by MIS Coordinator, approved by Admin (Sangeeta Shah) or Super Admin (Ram Shah)."
    ws9["A20"].font = Font(name="Calibri", size=10, color="666666")
    auto_width(ws9)

    # ── Sheet 10: Clawback Rules ──
    ws10 = wb.create_sheet("Clawback Rules")
    ws10.merge_cells("A1:D1")
    ws10["A1"] = "SIP & Lumpsum Clawback Rules"
    ws10["A1"].font = title_font

    cb_headers = ["Scenario", "Action", "Clawback", "Trail Impact"]
    for i, h in enumerate(cb_headers, 1):
        ws10.cell(row=3, column=i, value=h)
    style_header_row(ws10, 3, 4)

    cb_data = [
        ["SIP stopped by client", "Clawback upfront incentive paid on SIP", "Yes — proportional", "No trail freeze"],
        ["SIP redeemed early", "Clawback upfront incentive paid on SIP", "Yes — proportional", "No trail freeze"],
        ["Lumpsum redeemed early", "Clawback upfront incentive", "Yes — proportional", "No trail freeze"],
        ["SIP continues normally", "No action", "None", "Trail continues"],
        ["Switch within AMC", "No clawback", "None", "Trail continues on new scheme"],
    ]
    for r, row_data in enumerate(cb_data, 4):
        for c, val in enumerate(row_data, 1):
            ws10.cell(row=r, column=c, value=val)
        style_data_row(ws10, r, 4, alt=(r % 2 == 0))

    ws10.merge_cells("A10:D10")
    ws10["A10"] = "Note: Trail freeze is NOT needed. Only upfront incentive is clawed back."
    ws10["A10"].font = Font(name="Calibri", bold=True, size=10, color=RED)
    auto_width(ws10)

    # ── Sheet 11: Approval Workflow ──
    ws11 = wb.create_sheet("Approval Workflow")
    ws11.merge_cells("A1:D1")
    ws11["A1"] = "6-Step Payout Approval Workflow"
    ws11["A1"].font = title_font

    wf_headers = ["Step", "Role", "Action", "System Status"]
    for i, h in enumerate(wf_headers, 1):
        ws11.cell(row=3, column=i, value=h)
    style_header_row(ws11, 3, 4)

    wf_data = [
        ["1", "VJ Infosoft (External)", "Uploads raw Excel/CSV data", "Data Uploaded"],
        ["2", "Maker (MIS Coordinator)", "Fills remaining required fields, validates data", "Pending Checker Review"],
        ["3", "Checker", "Reviews & approves maker's entries", "Pending LOB Approval"],
        ["4", "LOB Heads (Life / Health / GI)", "Each LOB head approves their line of business", "Pending Finance Review"],
        ["5", "Finance Team Head", "Verifies amounts, confirms budget, approves for payment", "Pending Final Approval"],
        ["6", "Admin (Sangeeta Shah) / Super Admin (Ram Shah)", "Final payment approval and release", "Approved — Pay"],
    ]
    for r, row_data in enumerate(wf_data, 4):
        for c, val in enumerate(row_data, 1):
            ws11.cell(row=r, column=c, value=val)
        style_data_row(ws11, r, 4, alt=(r % 2 == 0))

    ws11.merge_cells("A11:D11")
    ws11["A11"] = "MF Trail Data: Entered by MIS Coordinator → Approved by Admin/Super Admin (separate flow)"
    ws11["A11"].font = Font(name="Calibri", bold=True, size=10, color=NAVY)

    ws11.merge_cells("A13:D13")
    ws11["A13"] = "Cost recovery minimum is for admin visibility only — NOT shown in employee logins."
    ws11["A13"].font = Font(name="Calibri", italic=True, size=10, color=RED)
    auto_width(ws11)

    # ── Sheet 12: POSP Naming ──
    ws12 = wb.create_sheet("POSP Naming")
    ws12.merge_cells("A1:D1")
    ws12["A1"] = "POSP Coded Naming Convention"
    ws12["A1"].font = title_font

    name_headers = ["Component", "Format", "Example", "Purpose"]
    for i, h in enumerate(name_headers, 1):
        ws12.cell(row=3, column=i, value=h)
    style_header_row(ws12, 3, 4)

    name_data = [
        ["Prefix", "TIB", "TIB", "Trustner Insurance Broking"],
        ["Channel Type", "POSP / BQP / FRAN", "POSP", "Point of Sales Person / BQP / Franchise"],
        ["Serial Number", "XXX (3+ digits)", "001", "Sequential unique ID"],
        ["Financial Year", "YYZZ", "2526", "FY 2025-26 onboarding year"],
        ["Full Code", "TIB/TYPE/XXX/YYZZ", "TIB/POSP/001/2526", "Complete coded identifier"],
    ]
    for r, row_data in enumerate(name_data, 4):
        for c, val in enumerate(row_data, 1):
            ws12.cell(row=r, column=c, value=val)
        style_data_row(ws12, r, 4, alt=(r % 2 == 0))

    ws12.merge_cells("A10:D10")
    ws12["A10"] = "Purpose: Information control — POSPs see their code, not others'. Prevents poaching."
    ws12["A10"].font = Font(name="Calibri", size=10, color="666666")
    auto_width(ws12)

    # ── Sheet 13: Removals ──
    ws13 = wb.create_sheet("Q1 FY27 Removals")
    ws13.merge_cells("A1:C1")
    ws13["A1"] = "Key Removals / Changes for Q1 FY27"
    ws13["A1"].font = title_font

    rem_headers = ["Item Removed", "Reason", "Future Status"]
    for i, h in enumerate(rem_headers, 1):
        ws13.cell(row=3, column=i, value=h)
    style_header_row(ws13, 3, 3)

    rem_data = [
        ["Multiplier concept", "Replaced with flat slab-based incentive", "Permanent removal"],
        ["CLV (Client Lifetime Value) tier bonuses", "Simplification of incentive structure", "Permanent removal"],
        ["FP Route approval for incentives", "Not needed for Q1 FY27", "May revisit in Q2 FY27"],
        ["Month freeze option", "Not needed in current workflow", "Permanent removal"],
        ["Management column in trail income", "Simplified to RM + Company split only", "Permanent removal"],
        ['"Credit" terminology', 'Renamed to "Weightage" throughout system', "Permanent change"],
        ["Area Manager / Regional Head slabs", "Deferred — to be decided", "Q2 FY27"],
    ]
    for r, row_data in enumerate(rem_data, 4):
        for c, val in enumerate(row_data, 1):
            ws13.cell(row=r, column=c, value=val)
        style_data_row(ws13, r, 3, alt=(r % 2 == 0))
    auto_width(ws13)

    # ── Sheet 14: Partner Hierarchy ──
    ws14 = wb.create_sheet("Partner Hierarchy")
    ws14.merge_cells("A1:F1")
    ws14["A1"] = "Partner Hierarchy & Channel Types"
    ws14["A1"].font = title_font
    ws14.merge_cells("A2:F2")
    ws14["A2"] = "Defines every partner type, their code format, recruiting ability, and default payout."
    ws14["A2"].font = Font(name="Calibri", italic=True, size=10, color="666666")

    ph_headers = ["Type", "Code Format", "Description", "Can Recruit POSPs", "Default Payout"]
    for i, h in enumerate(ph_headers, 1):
        ws14.cell(row=4, column=i, value=h)
    style_header_row(ws14, 4, 5)

    ph_data = [
        ["POSP", "TIB/POSP/XXX/YYZZ", "Point of Sales Person", "No", "Per category (A:50% to F2:90%)"],
        ["BQP", "TIB/BQP/XXX/YYZZ", "Business Quality Partner", "Yes (certified)", "85%"],
        ["Franchise (Basic)", "TIB/FRAN/XXX/YYZZ", "Franchise partner \u2014 self-sourcing only", "No", "80% (default)"],
        ["Franchise + BQP", "TIB/FRAN/XXX/YYZZ", "BQP-certified franchise \u2014 can recruit POSPs", "Yes", "80% default, up to 90% (Admin)"],
        ["Referral", "TIB/REF/XXX/YYZZ", "Refers business, earns fee", "No", "60% (Category C default)"],
    ]
    for r, row_data in enumerate(ph_data, 5):
        for c, val in enumerate(row_data, 1):
            ws14.cell(row=r, column=c, value=val)
        style_data_row(ws14, r, 5, alt=(r % 2 == 0))

    # Hierarchy Tracking Fields
    ws14.merge_cells("A11:F11")
    ws14["A11"] = "Hierarchy Tracking Fields"
    ws14["A11"].font = section_font

    ht_headers = ["Field", "Purpose"]
    for i, h in enumerate(ht_headers, 1):
        ws14.cell(row=13, column=i, value=h)
    style_header_row(ws14, 13, 2)

    ht_data = [
        ["managedByType", "DST / CDM / BQP / Franchise"],
        ["managedById", "ID of the managing entity"],
        ["parentFranchiseId", "Which Franchise the POSP falls under (if any)"],
        ["agreementPct", "Franchise/Referral overall agreement %"],
    ]
    for r, row_data in enumerate(ht_data, 14):
        for c, val in enumerate(row_data, 1):
            ws14.cell(row=r, column=c, value=val)
        style_data_row(ws14, r, 2, alt=(r % 2 == 0))

    # Upgrade Paths
    ws14.merge_cells("A19:F19")
    ws14["A19"] = "Upgrade Paths"
    ws14["A19"].font = section_font
    upgrade_paths = [
        "Referral \u2192 POSP (after POSP exam/certification)",
        "Referral \u2192 BQP (after BQP certification)",
        "Franchise (Basic) \u2192 Franchise + BQP (after BQP certification)",
    ]
    for i, path in enumerate(upgrade_paths):
        ws14[f"A{21+i}"] = path
        ws14[f"A{21+i}"].font = Font(name="Calibri", size=11)

    # POSP Category Approval Hierarchy
    ws14.merge_cells("A25:F25")
    ws14["A25"] = "POSP Category Approval Hierarchy"
    ws14["A25"].font = section_font

    appr_headers = ["Approver", "Can Approve Up To", "Categories", "Max Payout"]
    for i, h in enumerate(appr_headers, 1):
        ws14.cell(row=27, column=i, value=h)
    style_header_row(ws14, 27, 4)

    appr_data = [
        ["CDM", "Category C", "A, B, C", "Up to 60%"],
        ["Regional Head", "Category D+", "D, D+", "Up to 70%"],
        ["CDO", "Category E+", "E, E+", "Up to 80%"],
        ["Admin / Super Admin", "Category F2", "F1, F2", "Up to 90%"],
    ]
    for r, row_data in enumerate(appr_data, 28):
        for c, val in enumerate(row_data, 1):
            ws14.cell(row=r, column=c, value=val)
        style_data_row(ws14, r, 4, alt=(r % 2 == 0))

    ws14.merge_cells("A33:F33")
    ws14["A33"] = "Referral Payout Approval: Admin / Super Admin only (any change from default 60%)"
    ws14["A33"].font = Font(name="Calibri", bold=True, size=10, color=RED)
    auto_width(ws14)

    # ── Sheet 15: Franchise Payout ──
    ws15 = wb.create_sheet("Franchise Payout")
    ws15.merge_cells("A1:G1")
    ws15["A1"] = "Franchise Differential Payout Model"
    ws15["A1"].font = title_font
    ws15.merge_cells("A2:G2")
    ws15["A2"] = "Franchise earns the DIFFERENCE between their agreement % and the POSP category % under them."
    ws15["A2"].font = Font(name="Calibri", italic=True, size=10, color="666666")

    ws15.merge_cells("A4:G4")
    ws15["A4"] = "Formula: Franchise Payout = (Franchise Agreement % \u2212 POSP Category %) \u00d7 Commission"
    ws15["A4"].font = Font(name="Calibri", bold=True, size=11, color=GREEN)

    fp_headers = ["Scenario", "Commission", "POSP Category", "POSP Gets", "Franchise Agreement", "Franchise Gets", "Trustner Retains"]
    for i, h in enumerate(fp_headers, 1):
        ws15.cell(row=6, column=i, value=h)
    style_header_row(ws15, 6, 7)

    fp_data = [
        ["POSP D+ under 90% Franchise", "\u20b910,000", "D+ (70%)", "\u20b97,000", "90%", "\u20b92,000 (20%)", "\u20b91,000 (10%)"],
        ["POSP C under 85% Franchise", "\u20b910,000", "C (60%)", "\u20b96,000", "85%", "\u20b92,500 (25%)", "\u20b91,500 (15%)"],
        ["POSP A under 80% Franchise", "\u20b910,000", "A (50%)", "\u20b95,000", "80%", "\u20b93,000 (30%)", "\u20b92,000 (20%)"],
        ["Self-sourced by Franchise", "\u20b910,000", "N/A", "N/A", "80%", "\u20b98,000 (80%)", "\u20b92,000 (20%)"],
    ]
    for r, row_data in enumerate(fp_data, 7):
        for c, val in enumerate(row_data, 1):
            ws15.cell(row=r, column=c, value=val)
        style_data_row(ws15, r, 7, alt=(r % 2 == 0))

    # Franchise Agreement Range
    ws15.merge_cells("A12:G12")
    ws15["A12"] = "Franchise Agreement Range"
    ws15["A12"].font = section_font

    fr_headers = ["Default", "Range", "Who Can Change", "Notes"]
    for i, h in enumerate(fr_headers, 1):
        ws15.cell(row=14, column=i, value=h)
    style_header_row(ws15, 14, 4)

    fr_data = [
        ["80%", "80% to 90%", "Admin / Super Admin", "80% = mainly self-sourcing, Tier 2/3 cities"],
        ["Exception", "Above 90%", "Admin only", "Exceptional cases only"],
    ]
    for r, row_data in enumerate(fr_data, 15):
        for c, val in enumerate(row_data, 1):
            ws15.cell(row=r, column=c, value=val)
        style_data_row(ws15, r, 4, alt=(r % 2 == 0))

    ws15.merge_cells("A18:G18")
    ws15["A18"] = ("Business Rule: Franchise at 80% should NOT recruit higher-paying POSPs "
                    "(e.g., E+ or above) as the differential becomes too thin. "
                    "This works best in Tier 2/3 cities where POSP categories are typically lower (A-D).")
    ws15["A18"].font = Font(name="Calibri", bold=True, size=10, color=RED)
    auto_width(ws15)

    # ── Sheet 16: Referral System ──
    ws16 = wb.create_sheet("Referral System")
    ws16.merge_cells("A1:D1")
    ws16["A1"] = "Referral Code System"
    ws16["A1"].font = title_font

    ref_headers = ["Item", "Detail"]
    for i, h in enumerate(ref_headers, 1):
        ws16.cell(row=3, column=i, value=h)
    style_header_row(ws16, 3, 2)

    ref_data = [
        ["Code Format", "TIB/REF/001/2526"],
        ["Default Payout", "60% (= Category C equivalent)"],
        ["Rate Type", "Single flat rate across all LOBs"],
        ["Payout Change", "Admin / Super Admin approval required"],
        ["Can Recruit POSPs", "No"],
        ["Upgrade Path", "Can be upgraded to POSP or BQP with certification"],
    ]
    for r, row_data in enumerate(ref_data, 4):
        for c, val in enumerate(row_data, 1):
            ws16.cell(row=r, column=c, value=val)
        style_data_row(ws16, r, 2, alt=(r % 2 == 0))
    auto_width(ws16)

    # Save
    xlsx_path = os.path.join(OUT_DIR, f"Trustner_Incentive_Reference_{DOC_DATE_SHORT}.xlsx")
    wb.save(xlsx_path)
    print(f"✅ Excel saved: {xlsx_path}")
    return xlsx_path


# ════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT
# ════════════════════════════════════════════════════════════════════
def build_word():
    doc = Document()

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ── Cover Page ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nTRUSTNER ASSET SERVICES PVT LTD")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("INCENTIVE & COMPENSATION FRAMEWORK")
    run2.font.size = Pt(18)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0xD4, 0xA8, 0x43)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"\nComprehensive Reference Document\n{DOC_DATE}")
    run3.font.size = Pt(14)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("\nARN-286886 | IRDAI Corporate Agent (Composite) — Applied\nConfidential — Internal Use Only")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ── Table of Contents ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. DST (Direct Sales Team) Incentive Slabs",
        "2. POSP RM / RM POSP / RM Wealth Incentive Slabs",
        "3. Product Tier & Weightage System",
        "4. 10-Step Incentive Calculation Formula",
        "5. Channel Payout Structure",
        "6. POSP Category System (4 LOBs × 9 Categories)",
        "7. Motor Insurance Commission Framework",
        "8. Trail Income Distribution",
        "9. SIP & Lumpsum Clawback Rules",
        "10. MF Data Approval Flow",
        "11. 6-Step Payout Approval Workflow",
        "12. POSP Coded Naming Convention",
        "13. Area Manager / Regional Head (Deferred to Q2 FY27)",
        "14. Key Removals for Q1 FY27",
        "15. Partner Hierarchy & Channel Types",
        "16. Franchise Differential Payout Model",
        "17. Referral Code System",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ── Helper for tables ──
    def add_styled_table(headers, data, col_widths=None):
        table = doc.add_table(rows=1 + len(data), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        # Data rows
        for r, row_data in enumerate(data):
            for c, val in enumerate(row_data):
                cell = table.rows[r + 1].cells[c]
                cell.text = str(val)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        return table

    # ── 1. DST Slabs ──
    doc.add_heading("1. DST (Direct Sales Team) Incentive Slabs", level=1)
    doc.add_paragraph(
        "The DST incentive is a flat slab-based system with no multiplier. "
        "The incentive percentage is applied directly to the weighted business amount."
    )
    add_styled_table(
        ["Slab", "Achievement Range", "Incentive %", "Tier Name"],
        [
            ["1", "0 – 80%", "0%", "No Incentive"],
            ["2", "80.01 – 100%", "4%", "Base"],
            ["3", "100.01 – 125%", "5%", "Enhanced"],
            ["4", "125.01 – 150%", "6%", "Super"],
            ["5", "150.01% +", "8%", "Champion"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Formula: ")
    run.font.bold = True
    p.add_run("Incentive = Weighted Business × Slab %")
    doc.add_paragraph(
        "Example: A DST member achieves 112% of target with weighted business of ₹15,00,000. "
        "Slab = Enhanced (5%). Incentive = ₹15,00,000 × 5% = ₹75,000."
    )

    # ── 2. POSP RM Slabs ──
    doc.add_heading("2. POSP RM / RM POSP / RM Wealth Incentive Slabs", level=1)
    doc.add_paragraph(
        "Applicable to all RM roles managing POSP, sub-broker, and franchise channels. "
        "Same slab structure as DST but with lower percentages reflecting the management role."
    )
    add_styled_table(
        ["Slab", "Achievement Range", "Incentive %", "Tier Name"],
        [
            ["1", "0 – 80%", "0%", "No Incentive"],
            ["2", "80.01 – 100%", "1.2%", "Base"],
            ["3", "100.01 – 125%", "1.5%", "Enhanced"],
            ["4", "125.01 – 150%", "1.8%", "Super"],
            ["5", "150.01% +", "2.4%", "Champion"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Formula: ")
    run.font.bold = True
    p.add_run("Incentive = Team's Weighted Business × Slab %")

    # ── 3. Product Tiers ──
    doc.add_heading("3. Product Tier & Weightage System", level=1)
    doc.add_paragraph(
        'Products are classified into 4 tiers based on the commission percentage earned. '
        'The tier determines the "weightage" (previously called "credit") — i.e., how much of '
        'the premium counts toward the employee\'s target achievement.'
    )
    add_styled_table(
        ["Tier", "Commission Earned", "Weightage %", "Example Products"],
        [
            ["Tier 1", "≥ 30%", "100%", "Traditional Life, High-Commission Health"],
            ["Tier 2", "15% – 30%", "75%", "ULIP, Standard Health, Fire"],
            ["Tier 3", "< 15%", "50%", "Term Life, Motor, Standard GI"],
            ["Tier 4", "≤ 7.5%", "25%", "Group Term, Group Mediclaim, Accidental"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Weighted Business = ")
    run.font.bold = True
    p.add_run("Σ (Each Policy Premium × Tier Weightage %)")
    doc.add_paragraph(
        "Example: ₹1,00,000 Tier 1 policy (100%) + ₹2,00,000 Tier 3 policy (50%) "
        "= ₹1,00,000 + ₹1,00,000 = ₹2,00,000 total weighted business."
    )

    # ── 4. 10-Step Formula ──
    doc.add_heading("4. 10-Step Incentive Calculation Pipeline", level=1)
    add_styled_table(
        ["Step", "Action", "Formula / Logic", "Output"],
        [
            ["1", "Collect Business Entries", "All policies in the month", "Raw premium list"],
            ["2", "Assign Product Tier", "Based on commission %", "Tier 1/2/3/4 per policy"],
            ["3", "Calculate Weighted Business", "Premium × tier weightage %", "Weighted premium per policy"],
            ["4", "Sum Weighted Business", "Σ all weighted premiums", "Total weighted business (₹)"],
            ["5", "Get Monthly Target", "From employee master / goal sheet", "Target amount (₹)"],
            ["6", "Calculate Achievement %", "(Weighted ÷ Target) × 100", "Achievement %"],
            ["7", "Look Up Slab", "DST or POSP RM slab table", "Incentive %"],
            ["8", "Calculate Gross Incentive", "Weighted Business × Slab %", "Gross incentive (₹)"],
            ["9", "Apply Deductions", "Clawbacks, recoveries", "Net incentive (₹)"],
            ["10", "Route for Approval", "6-step workflow", "Final approved payout"],
        ]
    )

    # ── 5. Channel Payouts ──
    doc.add_heading("5. Channel Payout Structure", level=1)
    doc.add_paragraph(
        "Defines how much of the company's commission is shared with each distribution channel."
    )
    add_styled_table(
        ["Channel", "Payout to Channel", "Company Margin", "Notes"],
        [
            ["POSP", "70%", "30%", "Standard POSP agents"],
            ["BQP (Referral)", "85%", "15%", "Revised from 65% to 85%"],
            ["Referral", "50%", "50%", "Casual referral partners"],
            ["Employee (DST)", "0% (salaried)", "100%", "Incentive via DST slabs"],
            ["Sub-Broker", "60%", "40%", "Registered sub-brokers"],
            ["Franchise", "80% – 90%", "10% – 20%", "Range based on agreement"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Key Changes: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run("BQP revised from 65% to 85%. Franchise is now a range (80-90%), not flat 85%.")

    # ── 6. POSP Categories ──
    doc.add_heading("6. POSP Category System — 4 LOBs × 9 Categories", level=1)
    doc.add_paragraph(
        "Each POSP is assigned a category independently for each of the 4 Lines of Business (LOBs). "
        "The category determines the percentage of the company's commission that flows to the POSP. "
        "There are 9 categories ranging from A (50%) to F2 (90%)."
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("4 LOBs: ")
    run.font.bold = True
    p.add_run("Life Insurance | Health Insurance | GI Non-Motor | GI Motor")

    add_styled_table(
        ["Category", "Payout %", "Description"],
        [
            ["A", "50%", "Entry level / new POSP"],
            ["B", "55%", "Developing"],
            ["C", "60%", "Established"],
            ["D", "65%", "Performing"],
            ["D+", "70%", "High performing"],
            ["E", "75%", "Senior"],
            ["E+", "80%", "Senior+"],
            ["F1", "85%", "Star performer"],
            ["F2", "90%", "Top performer — maximum payout"],
        ]
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "A POSP can hold different categories for different LOBs. For example: "
        "Category C in Life, Category E in Health, Category B in GI Non-Motor, Category D in GI Motor."
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Category Review: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(
        "Category assignment is reviewed manually at the end of each quarter by "
        "Admin (Sangeeta Shah) / Super Admin (Ram Shah). No automatic progression."
    )

    # ── 7. Motor Insurance ──
    doc.add_heading("7. Motor Insurance Commission Framework", level=1)
    doc.add_paragraph(
        "Motor insurance has a unique commission structure due to the OD/TP split. "
        "Commission is earned only on the OD (Own Damage) portion. "
        "TP (Third Party) premium rates are fixed by IRDAI and carry no commission."
    )
    add_styled_table(
        ["Component", "Commission", "Notes"],
        [
            ["OD Premium", "10% – 15%", "Commission earned by Trustner from insurer"],
            ["TP Premium", "0%", "No commission — IRDAI-regulated rate"],
            ["POSP Share", "Per category %", "Category A(50%) to F2(90%) of Trustner's OD commission"],
            ["BQP Share", "85%", "Of Trustner's OD commission"],
            ["Franchise Share", "80-90%", "Range per agreement"],
        ]
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Product Tier: Motor falls under Tier 3 (< 15% commission) → 50% weightage for incentive calculation."
    )
    doc.add_paragraph(
        "Example: ₹50,000 motor policy (OD ₹30,000 + TP ₹20,000). "
        "Commission on OD = 15% = ₹4,500. "
        "If POSP is Category D (65%): POSP gets ₹2,925, Trustner retains ₹1,575."
    )

    # ── 8. Trail Income ──
    doc.add_heading("8. Trail Income Distribution — Mutual Funds", level=1)
    doc.add_paragraph(
        "Trail income from mutual fund AUM is split between the RM and the company. "
        "The management column has been removed — it is now a direct 2-way split."
    )
    add_styled_table(
        ["Source Type", "RM Share", "Company Share", "Condition"],
        [
            ["Self-Sourced — New PAN", "25%", "75%", "RM brings completely new client"],
            ["Self-Sourced — Existing PAN", "15%", "85%", "RM brings new biz from existing client"],
            ["Assigned — New Business", "15%", "85%", "Company assigns client, RM generates new biz"],
            ["Assigned — No New Business", "0%", "100%", "Company assigns client, no new business"],
            ["Office Walk-in", "10%", "90%", "Client walks in, RM services them"],
        ]
    )
    doc.add_paragraph("")
    doc.add_heading("Trail Rate Reference (AMC to Trustner ARN-286886)", level=2)
    add_styled_table(
        ["Fund Category", "Typical Trail Rate"],
        [
            ["Equity — Active", "0.70% – 1.50%"],
            ["Equity — Index / Passive", "0.30% – 0.50%"],
            ["Hybrid / Balanced", "0.50% – 1.00%"],
            ["Debt — Active", "0.20% – 0.80%"],
            ["Liquid / Overnight", "0.05% – 0.15%"],
        ]
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "MF data is entered by the MIS Coordinator and approved by Admin (Sangeeta Shah) or Super Admin (Ram Shah)."
    )

    # ── 9. Clawback ──
    doc.add_heading("9. SIP & Lumpsum Clawback Rules", level=1)
    doc.add_paragraph(
        "Upfront incentive paid on SIP/Lumpsum business is subject to clawback if the client "
        "stops the SIP or redeems the lumpsum investment. Trail freeze is NOT applied."
    )
    add_styled_table(
        ["Scenario", "Clawback", "Trail Impact"],
        [
            ["SIP stopped by client", "Yes — proportional clawback of upfront incentive", "No trail freeze"],
            ["SIP redeemed early", "Yes — proportional clawback", "No trail freeze"],
            ["Lumpsum redeemed early", "Yes — proportional clawback", "No trail freeze"],
            ["SIP continues normally", "None", "Trail continues"],
            ["Switch within AMC", "None", "Trail continues on new scheme"],
        ]
    )

    # ── 10. MF Data Approval ──
    doc.add_heading("10. MF Data Approval Flow", level=1)
    doc.add_paragraph(
        "Mutual fund business data follows a simplified 2-step approval:"
    )
    doc.add_paragraph("Step 1: MIS Coordinator enters/uploads MF transaction data", style='List Bullet')
    doc.add_paragraph("Step 2: Admin (Sangeeta Shah) or Super Admin (Ram Shah) reviews and approves", style='List Bullet')
    doc.add_paragraph("")
    doc.add_paragraph(
        "This is separate from the insurance payout workflow. MF trail distribution follows "
        "the trail income table in Section 8."
    )

    # ── 11. Payout Workflow ──
    doc.add_heading("11. 6-Step Payout Approval Workflow", level=1)
    doc.add_paragraph(
        "All insurance incentive payouts go through a 6-step maker-checker-approver workflow:"
    )
    add_styled_table(
        ["Step", "Role", "Action"],
        [
            ["1", "VJ Infosoft (External)", "Uploads raw Excel/CSV business data"],
            ["2", "Maker (MIS Coordinator)", "Fills remaining fields, validates entries"],
            ["3", "Checker", "Reviews and approves maker's entries"],
            ["4", "LOB Heads (Life / Health / GI)", "Each LOB head approves their line of business"],
            ["5", "Finance Team Head", "Verifies amounts, confirms budget, approves for payment"],
            ["6", "Admin / Super Admin", "Final payment approval and release"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Important: ")
    run.font.bold = True
    p.add_run(
        "Cost recovery minimum is for admin visibility only — it is NOT shown in employee logins."
    )

    # ── 12. POSP Naming ──
    doc.add_heading("12. POSP Coded Naming Convention", level=1)
    doc.add_paragraph(
        "All POSPs, BQPs, and Franchise partners are assigned a coded identifier "
        "for information control. This prevents poaching and maintains confidentiality."
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Format: ")
    run.font.bold = True
    p.add_run("TIB / CHANNEL TYPE / SERIAL / FINANCIAL YEAR")
    doc.add_paragraph("")
    add_styled_table(
        ["Component", "Format", "Example"],
        [
            ["Prefix", "TIB", "Trustner Insurance Broking"],
            ["Channel Type", "POSP / BQP / FRAN", "Agent type"],
            ["Serial Number", "XXX (3+ digits)", "001, 002, ..."],
            ["Financial Year", "YYZZ", "2526 = FY 2025-26"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Example: ")
    run.font.bold = True
    p.add_run("TIB/POSP/001/2526 — First POSP onboarded in FY 2025-26")

    # ── 13. Area Manager ──
    doc.add_heading("13. Area Manager / Regional Head", level=1)
    doc.add_paragraph(
        "Incentive slabs for Area Manager and Regional Head roles will be decided and "
        "implemented from Quarter 2 of FY 2026-27 (July 2026 onwards). "
        "No slabs are active for Q1 FY27."
    )

    # ── 14. Removals ──
    doc.add_heading("14. Key Removals & Changes for Q1 FY27", level=1)
    add_styled_table(
        ["Item", "Change", "Status"],
        [
            ["Multiplier concept", "Replaced with flat slab-based incentive", "Permanent"],
            ["CLV tier bonuses", "Removed — simplification", "Permanent"],
            ["FP Route approval", "Not needed for Q1 FY27", "May revisit Q2"],
            ["Month freeze option", "Removed", "Permanent"],
            ["Management column (trail)", "Removed — RM + Company only", "Permanent"],
            ['"Credit" terminology', 'Renamed to "Weightage"', "Permanent"],
            ["AM / RH slabs", "Deferred", "Q2 FY27"],
        ]
    )

    # ── 15. Partner Hierarchy & Channel Types ──
    doc.add_heading("15. Partner Hierarchy & Channel Types", level=1)
    doc.add_paragraph(
        "Defines every partner type within Trustner's distribution network, "
        "their coded identifier format, recruiting capabilities, and default payout levels."
    )
    add_styled_table(
        ["Type", "Code Format", "Description", "Can Recruit POSPs", "Default Payout"],
        [
            ["POSP", "TIB/POSP/XXX/YYZZ", "Point of Sales Person", "No", "Per category (A:50% to F2:90%)"],
            ["BQP", "TIB/BQP/XXX/YYZZ", "Business Quality Partner", "Yes (certified)", "85%"],
            ["Franchise (Basic)", "TIB/FRAN/XXX/YYZZ", "Franchise partner \u2014 self-sourcing only", "No", "80% (default)"],
            ["Franchise + BQP", "TIB/FRAN/XXX/YYZZ", "BQP-certified franchise \u2014 can recruit POSPs", "Yes", "80% default, up to 90% (Admin)"],
            ["Referral", "TIB/REF/XXX/YYZZ", "Refers business, earns fee", "No", "60% (Category C default)"],
        ]
    )

    doc.add_paragraph("")
    doc.add_heading("Hierarchy Tracking Fields", level=2)
    add_styled_table(
        ["Field", "Purpose"],
        [
            ["managedByType", "DST / CDM / BQP / Franchise"],
            ["managedById", "ID of the managing entity"],
            ["parentFranchiseId", "Which Franchise the POSP falls under (if any)"],
            ["agreementPct", "Franchise/Referral overall agreement %"],
        ]
    )

    doc.add_paragraph("")
    doc.add_heading("Upgrade Paths", level=2)
    doc.add_paragraph("Referral \u2192 POSP (after POSP exam/certification)", style='List Bullet')
    doc.add_paragraph("Referral \u2192 BQP (after BQP certification)", style='List Bullet')
    doc.add_paragraph("Franchise (Basic) \u2192 Franchise + BQP (after BQP certification)", style='List Bullet')

    doc.add_paragraph("")
    doc.add_heading("POSP Category Approval Hierarchy", level=2)
    doc.add_paragraph(
        "Defines which role can approve POSP category assignments up to a given level."
    )
    add_styled_table(
        ["Approver", "Can Approve Up To", "Categories", "Max Payout"],
        [
            ["CDM", "Category C", "A, B, C", "Up to 60%"],
            ["Regional Head", "Category D+", "D, D+", "Up to 70%"],
            ["CDO", "Category E+", "E, E+", "Up to 80%"],
            ["Admin / Super Admin", "Category F2", "F1, F2", "Up to 90%"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Referral Payout Approval: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run("Admin / Super Admin only (any change from default 60%).")

    # ── 16. Franchise Differential Payout Model ──
    doc.add_heading("16. Franchise Differential Payout Model", level=1)
    doc.add_paragraph(
        "The franchise earns the DIFFERENCE between their agreement percentage and the "
        "POSP category percentage for POSPs under them. For self-sourced business, "
        "the franchise earns their full agreement percentage."
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Formula: ")
    run.font.bold = True
    p.add_run("Franchise Payout = (Franchise Agreement % \u2212 POSP Category %) \u00d7 Commission")

    doc.add_paragraph("")
    add_styled_table(
        ["Scenario", "Commission", "POSP Category", "POSP Gets", "Franchise Agreement", "Franchise Gets", "Trustner Retains"],
        [
            ["POSP D+ under 90% Franchise", "\u20b910,000", "D+ (70%)", "\u20b97,000", "90%", "\u20b92,000 (20%)", "\u20b91,000 (10%)"],
            ["POSP C under 85% Franchise", "\u20b910,000", "C (60%)", "\u20b96,000", "85%", "\u20b92,500 (25%)", "\u20b91,500 (15%)"],
            ["POSP A under 80% Franchise", "\u20b910,000", "A (50%)", "\u20b95,000", "80%", "\u20b93,000 (30%)", "\u20b92,000 (20%)"],
            ["Self-sourced by Franchise", "\u20b910,000", "N/A", "N/A", "80%", "\u20b98,000 (80%)", "\u20b92,000 (20%)"],
        ]
    )

    doc.add_paragraph("")
    doc.add_heading("Franchise Agreement Range", level=2)
    add_styled_table(
        ["Default", "Range", "Who Can Change", "Notes"],
        [
            ["80%", "80% to 90%", "Admin / Super Admin", "80% = mainly self-sourcing, Tier 2/3 cities"],
            ["Exception", "Above 90%", "Admin only", "Exceptional cases only"],
        ]
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Business Rule: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(
        "Franchise at 80% should NOT recruit higher-paying POSPs (e.g., E+ or above) "
        "as the differential becomes too thin. This works best in Tier 2/3 cities where "
        "POSP categories are typically lower (A-D)."
    )

    # ── 17. Referral Code System ──
    doc.add_heading("17. Referral Code System", level=1)
    doc.add_paragraph(
        "The referral channel provides a simple, low-commitment way for individuals to "
        "earn a fee by referring insurance or mutual fund business to Trustner."
    )
    add_styled_table(
        ["Item", "Detail"],
        [
            ["Code Format", "TIB/REF/001/2526"],
            ["Default Payout", "60% (= Category C equivalent)"],
            ["Rate Type", "Single flat rate across all LOBs"],
            ["Payout Change", "Admin / Super Admin approval required"],
            ["Can Recruit POSPs", "No"],
            ["Upgrade Path", "Can be upgraded to POSP or BQP with certification"],
        ]
    )

    # ── Footer ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\nDocument prepared on {DOC_DATE}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Trustner Asset Services Pvt Ltd").font.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Confidential — For internal reference only")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    docx_path = os.path.join(OUT_DIR, f"Trustner_Incentive_Reference_{DOC_DATE_SHORT}.docx")
    doc.save(docx_path)
    print(f"✅ Word saved: {docx_path}")
    return docx_path


# ════════════════════════════════════════════════════════════════════
#  MAIN
# ════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print(f"Generating Trustner Incentive Reference Documents — {DOC_DATE}")
    print("=" * 60)
    xlsx = build_excel()
    docx = build_word()
    print("=" * 60)
    print("Done! Files created:")
    print(f"  📊 {xlsx}")
    print(f"  📄 {docx}")
